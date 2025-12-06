"""
Hàm xuất Word cho Bảo Vệ An Ninh
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_bao_ve_an_ninh(personnel_list: List[Personnel],
                                don_vi: str = "Đại đội 3",
                                tieu_doan: str = "TIỂU ĐOÀN 38",
                                dia_diem: str = "Đắk Lắk",
                                nam: str = "2025",
                                ngay_bo_sung: str = "01",
                                thang_bo_sung: str = "7",
                                nam_bo_sung: str = "2025",
                                chinh_tri_vien: str = "",
                                db_service=None) -> bytes:
    """
    Xuất danh sách Bảo Vệ An Ninh ra file Word với format giống mẫu
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang Letter (27.94 cm x 21.59 cm)
        section = doc.sections[0]
        section.page_width = Inches(11.0)  # 27.94 cm = 11 inches
        section.page_height = Inches(8.5)  # 21.59 cm = 8.5 inches
        section.left_margin = Inches(0.79)  # 2 cm
        section.right_margin = Inches(0.79)  # 2 cm
        section.top_margin = Inches(0.79)  # 2 cm
        section.bottom_margin = Inches(0.79)  # 2 cm
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: TIỂU ĐOÀN 38, ĐẠI ĐỘI 3
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_run1 = left_cell.paragraphs[0].add_run(tieu_doan)
        left_run1.font.size = Pt(11)
        left_run1.font.name = 'Times New Roman'
        left_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_cell.paragraphs[0].add_run("\n")
        left_run2 = left_cell.paragraphs[0].add_run(don_vi.upper())
        left_run2.font.size = Pt(11)
        left_run2.font.name = 'Times New Roman'
        left_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Cột phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM, Độc lập - Tự do - Hạnh phúc
        right_cell = header_table.rows[0].cells[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        right_run1 = right_cell.paragraphs[0].add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        right_run1.font.size = Pt(11)
        right_run1.font.name = 'Times New Roman'
        right_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run1.bold = True
        
        right_cell.paragraphs[0].add_run("\n")
        right_run2 = right_cell.paragraphs[0].add_run("Độc lập - Tự do - Hạnh phúc")
        right_run2.font.size = Pt(11)
        right_run2.font.name = 'Times New Roman'
        right_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run2.italic = True
        right_run2.underline = True
        
        # Merge cells cho dòng 2 (date line)
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
        date_cell = header_table.rows[1].cells[0]
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_cell.paragraphs[0].clear()
        
        date_text = f"{dia_diem}, ngày tháng năm {nam}"
        date_run = date_cell.paragraphs[0].add_run(date_text)
        date_run.font.size = Pt(11)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Bỏ borders cho header table
        for row in header_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tiêu đề: DANH SÁCH
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DANH SÁCH")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle: Bí thư cấp uỷ, chi bộ phụ trách công tác bảo vệ an ninh và chiến sỹ bảo vệ năm 2025
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Bí thư cấp uỷ, chi bộ phụ trách công tác bảo vệ an ninh và chiến sỹ bảo vệ năm {nam}")
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle 2: (kiện toàn, bổ sung ngày 01 tháng 7 năm 2025)
        subtitle2 = doc.add_paragraph()
        subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle2_run = subtitle2.add_run(f"(kiện toàn, bổ sung ngày {ngay_bo_sung} tháng {thang_bo_sung} năm {nam_bo_sung})")
        subtitle2_run.font.size = Pt(10)
        subtitle2_run.font.name = 'Times New Roman'
        subtitle2_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 11 cột
        table = doc.add_table(rows=1, cols=11)
        table.style = None  # Bỏ style mặc định
        
        # Thiết lập borders cho toàn bộ table
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Thêm borders cho toàn bộ table
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Độ rộng cột
        col_widths = [
            Inches(0.4),   # TT
            Inches(1.5),   # Họ và tên (Ngày, tháng, năm sinh)
            Inches(1.2),   # NN CB-CV Đ.VI
            Inches(1.0),   # Đăng đoàn
            Inches(0.8),   # DT TG
            Inches(0.5),   # VH
            Inches(2.5),   # Họ tên cha/mẹ/vợ
            Inches(2.0),   # Nguyên quán Trú quán
            Inches(0.8),   # Thời gian vào
            Inches(0.8),   # Thời gian ra
            Inches(0.8),   # Ghi chú
        ]
        
        for idx, width in enumerate(col_widths):
            if idx < len(table.columns):
                table.columns[idx].width = width
        
        # Header row
        header_cells = table.rows[0].cells
        headers = [
            'TT',
            'Họ và tên\n(Ngày, tháng,\nnăm sinh)',
            'NN\nCB-CV\nĐ.VI',
            'Đăng\ndoàn',
            'DT\nTG',
            'VH',
            'Họ tên cha\n(Năm sinh-nghề nghiệp) /\nHọ tên mẹ\n(Năm sinh-nghề nghiệp) /\nHọ tên vợ\n(Năm sinh-nghề nghiệp)',
            'Nguyên quán\nTrú quán',
            'Thời gian\nvào',
            'Thời gian\nra',
            'Ghi chú'
        ]
        
        for i, header_text in enumerate(headers):
            if i < len(header_cells):
                cell = header_cells[i]
                cell.text = header_text
                
                # Format header
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho header
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm dữ liệu
        for idx, p in enumerate(personnel_list, 1):
            row = table.add_row()
            cells = row.cells
            
            # Format tất cả cells
            for cell in cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho từng cell
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Cột 1: TT
            cells[0].text = str(idx)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cột 2: Họ và tên (Ngày, tháng, năm sinh)
            ho_ten = p.hoTen or ''
            ngay_sinh = p.ngaySinh or ''
            if ho_ten and ngay_sinh:
                cells[1].text = f"{ho_ten}\n({ngay_sinh})"
            elif ho_ten:
                cells[1].text = ho_ten
            else:
                cells[1].text = ''
            
            # Cột 3: NN CB-CV Đ.VI
            nhap_ngu = p.nhapNgu or ''
            # Format: MM/YYYY nếu có
            if '/' in nhap_ngu and len(nhap_ngu.split('/')) == 3:
                parts = nhap_ngu.split('/')
                nhap_ngu = f"{parts[1]}/{parts[2]}"
            
            cap_bac = p.capBac or ''
            chuc_vu = p.chucVu or ''
            don_vi = p.donVi or ''
            
            nn_cb_cv = []
            if nhap_ngu:
                nn_cb_cv.append(nhap_ngu)
            if cap_bac and chuc_vu:
                nn_cb_cv.append(f"{cap_bac}/{chuc_vu}")
            elif cap_bac:
                nn_cb_cv.append(cap_bac)
            elif chuc_vu:
                nn_cb_cv.append(chuc_vu)
            if don_vi:
                nn_cb_cv.append(don_vi)
            
            cells[2].text = "\n".join(nn_cb_cv) if nn_cb_cv else ''
            
            # Cột 4: Đăng đoàn
            dang_doan_parts = []
            if p.thongTinKhac and p.thongTinKhac.dang:
                ngay_vao_dang = p.thongTinKhac.dang.ngayVao or ''
                if ngay_vao_dang:
                    # Format: DD/MM/YYYY -> DD/MM/YYYY
                    if '/' in ngay_vao_dang:
                        parts = ngay_vao_dang.split('/')
                        if len(parts) == 3:
                            ngay_vao_dang = f"{parts[0]}/{parts[1]}/{parts[2]}"
                    dang_doan_parts.append(f"Đảng {ngay_vao_dang}")
            
            if p.thongTinKhac and p.thongTinKhac.doan:
                ngay_vao_doan = p.thongTinKhac.doan.ngayVao or ''
                if ngay_vao_doan:
                    if '/' in ngay_vao_doan:
                        parts = ngay_vao_doan.split('/')
                        if len(parts) == 3:
                            ngay_vao_doan = f"{parts[0]}/{parts[1]}/{parts[2]}"
                    dang_doan_parts.append(f"Đoàn {ngay_vao_doan}")
            
            cells[3].text = "\n".join(dang_doan_parts) if dang_doan_parts else ''
            
            # Cột 5: DT TG
            dan_toc = p.danToc or ''
            ton_giao = p.tonGiao or 'Không'
            cells[4].text = f"{dan_toc}\n{ton_giao}"
            
            # Cột 6: VH
            cells[5].text = p.trinhDoVanHoa or ''
            
            # Cột 7: Họ tên cha/mẹ/vợ
            gia_dinh_info = []
            bo_de = []
            me_de = []
            vo = []
            
            # Ưu tiên lấy từ bảng nguoi_than
            if db_service and p.id:
                try:
                    nguoi_than_list = db_service.get_nguoi_than_by_personnel(p.id)
                    
                    for nguoi_than in nguoi_than_list:
                        moi_quan_he = (nguoi_than.moiQuanHe or '').lower().strip()
                        ho_ten = (nguoi_than.hoTen or '').strip()
                        ngay_sinh = (nguoi_than.ngaySinh or '').strip()
                        noi_dung = (nguoi_than.noiDung or '').strip()
                        
                        if not ho_ten:
                            continue
                        
                        # Lấy năm sinh
                        nam_sinh = ""
                        if ngay_sinh:
                            try:
                                if '/' in ngay_sinh:
                                    parts = ngay_sinh.split('/')
                                    if len(parts) >= 3:
                                        nam_sinh = parts[-1]  # Lấy năm (phần cuối)
                                    elif len(parts) == 2:
                                        nam_sinh = parts[1]  # MM/YYYY
                                    else:
                                        nam_sinh = parts[0][:4] if len(parts[0]) >= 4 else parts[0]
                                elif '-' in ngay_sinh:
                                    parts = ngay_sinh.split('-')
                                    nam_sinh = parts[0][:4] if len(parts[0]) >= 4 else parts[0]
                                else:
                                    nam_sinh = ngay_sinh[:4] if len(ngay_sinh) >= 4 else ngay_sinh
                            except:
                                nam_sinh = ""
                        
                        # Nghề nghiệp
                        nghe = noi_dung if noi_dung else "làm nông"
                        
                        # Tạo chuỗi thông tin
                        if nam_sinh:
                            info_str = f"{ho_ten} ({nam_sinh}-{nghe})"
                        else:
                            info_str = f"{ho_ten} ({nghe})"
                        
                        # Phân loại theo quan hệ
                        # Bố đẻ / Cha đẻ
                        if ('bố đẻ' in moi_quan_he or 'cha đẻ' in moi_quan_he or 
                            (('bố' in moi_quan_he or 'cha' in moi_quan_he) and 
                             'vợ' not in moi_quan_he and 'vo' not in moi_quan_he and 
                             'đẻ' in moi_quan_he)):
                            bo_de.append(info_str)
                        # Mẹ đẻ
                        elif ('mẹ đẻ' in moi_quan_he or 'me đẻ' in moi_quan_he or 
                              (('mẹ' in moi_quan_he or 'me' in moi_quan_he) and 
                               'vợ' not in moi_quan_he and 'vo' not in moi_quan_he and 
                               'đẻ' in moi_quan_he)):
                            me_de.append(info_str)
                        # Vợ
                        elif 'vợ' in moi_quan_he or 'vo' in moi_quan_he:
                            vo.append(info_str)
                        # Bố (nếu không có "đẻ" thì coi là bố đẻ)
                        elif 'bố' in moi_quan_he or 'cha' in moi_quan_he:
                            if 'vợ' not in moi_quan_he and 'vo' not in moi_quan_he:
                                bo_de.append(info_str)
                        # Mẹ (nếu không có "đẻ" thì coi là mẹ đẻ)
                        elif 'mẹ' in moi_quan_he or 'me' in moi_quan_he:
                            if 'vợ' not in moi_quan_he and 'vo' not in moi_quan_he:
                                me_de.append(info_str)
                    
                except Exception as e:
                    # Log lỗi nhưng vẫn tiếp tục với fallback
                    pass
            
            # Fallback: sử dụng các field cũ nếu không có dữ liệu từ nguoi_than
            if not bo_de and p.hoTenCha:
                nam_sinh_cha = ""
                if p.ngaySinhCha:
                    try:
                        if '/' in p.ngaySinhCha:
                            parts = p.ngaySinhCha.split('/')
                            nam_sinh_cha = parts[-1] if len(parts) >= 3 else (parts[1] if len(parts) == 2 else parts[0][:4])
                        else:
                            nam_sinh_cha = p.ngaySinhCha[:4] if len(p.ngaySinhCha) >= 4 else p.ngaySinhCha
                    except:
                        pass
                bo_de.append(f"{p.hoTenCha} ({nam_sinh_cha}-làm nông)" if nam_sinh_cha else f"{p.hoTenCha} (làm nông)")
            
            if not me_de and p.hoTenMe:
                nam_sinh_me = ""
                if p.ngaySinhMe:
                    try:
                        if '/' in p.ngaySinhMe:
                            parts = p.ngaySinhMe.split('/')
                            nam_sinh_me = parts[-1] if len(parts) >= 3 else (parts[1] if len(parts) == 2 else parts[0][:4])
                        else:
                            nam_sinh_me = p.ngaySinhMe[:4] if len(p.ngaySinhMe) >= 4 else p.ngaySinhMe
                    except:
                        pass
                me_de.append(f"{p.hoTenMe} ({nam_sinh_me}-làm nông)" if nam_sinh_me else f"{p.hoTenMe} (làm nông)")
            
            if not vo and p.hoTenVo:
                nam_sinh_vo = ""
                if p.ngaySinhVo:
                    try:
                        if '/' in p.ngaySinhVo:
                            parts = p.ngaySinhVo.split('/')
                            nam_sinh_vo = parts[-1] if len(parts) >= 3 else (parts[1] if len(parts) == 2 else parts[0][:4])
                        else:
                            nam_sinh_vo = p.ngaySinhVo[:4] if len(p.ngaySinhVo) >= 4 else p.ngaySinhVo
                    except:
                        pass
                vo.append(f"{p.hoTenVo} ({nam_sinh_vo}-GV)" if nam_sinh_vo else f"{p.hoTenVo} (GV)")
            
            # Thêm vào danh sách theo thứ tự: Bố đẻ, Mẹ đẻ, Vợ
            gia_dinh_info.extend(bo_de)
            gia_dinh_info.extend(me_de)
            gia_dinh_info.extend(vo)
            
            cells[6].text = " / ".join(gia_dinh_info) if gia_dinh_info else ''
            
            # Cột 8: Nguyên quán Trú quán
            que_quan = p.queQuan or ''
            tru_quan = p.truQuan or ''
            if que_quan and tru_quan:
                cells[7].text = f"{que_quan} / {tru_quan}"
            elif que_quan:
                cells[7].text = que_quan
            elif tru_quan:
                cells[7].text = tru_quan
            else:
                cells[7].text = ''
            
            # Cột 9: Thời gian vào
            if db_service:
                bao_ve_info = db_service.get_bao_ve_an_ninh_info(p.id)
                cells[8].text = bao_ve_info.get('thoiGianVao', '') or ''
            else:
                cells[8].text = ''
            
            # Cột 10: Thời gian ra
            if db_service:
                bao_ve_info = db_service.get_bao_ve_an_ninh_info(p.id)
                cells[9].text = bao_ve_info.get('thoiGianRa', '') or ''
            else:
                cells[9].text = ''
            
            # Cột 11: Ghi chú
            cells[10].text = p.ghiChu or ''
        
        # Footer: CHÍNH TRỊ VIÊN, tên
        doc.add_paragraph()  # Khoảng trống
        
        footer_table = doc.add_table(rows=2, cols=1)
        footer_table.autofit = False
        footer_cell = footer_table.rows[0].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_cell.paragraphs[0].clear()
        footer_run1 = footer_cell.paragraphs[0].add_run("CHÍNH TRỊ VIÊN")
        footer_run1.font.size = Pt(11)
        footer_run1.font.name = 'Times New Roman'
        footer_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        footer_run1.bold = True
        
        footer_cell = footer_table.rows[1].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_cell.paragraphs[0].clear()
        if chinh_tri_vien:
            footer_run2 = footer_cell.paragraphs[0].add_run(chinh_tri_vien)
            footer_run2.font.size = Pt(11)
            footer_run2.font.name = 'Times New Roman'
            footer_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            footer_run2.bold = True
        
        # Bỏ borders cho footer table
        for row in footer_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("Cần cài đặt python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

