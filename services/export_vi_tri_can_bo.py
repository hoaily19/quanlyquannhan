"""
Hàm xuất Word cho Vị Trí Cán Bộ - tạm thời tách riêng để tránh conflict
"""
import io
import sys
from typing import List
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_vi_tri_can_bo(personnel_list: List[Personnel],
                               don_vi: str = "Đại đội 3",
                               nam: str = "2025",
                               chinh_tri_vien: str = "Đại úy Triệu Văn Dũng",
                               db_service=None) -> bytes:
    """
    Xuất danh sách Vị Trí Cán Bộ ra file Word với 11 cột giống mẫu
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang: A4 Landscape theo thông số ảnh
        # A4: 29.7 cm x 21 cm (Landscape)
        # Margins: Top 2 cm, Bottom 2 cm, Left 2.54 cm, Right 2.54 cm
        # Header/Footer distance: 1.25 cm
        section = doc.sections[0]
        section.page_width = Inches(11.69)  # 29.7 cm = 11.69 inches
        section.page_height = Inches(8.27)   # 21 cm = 8.27 inches
        section.left_margin = Inches(1.0)    # 2.54 cm = 1.0 inch
        section.right_margin = Inches(1.0)   # 2.54 cm = 1.0 inch
        section.top_margin = Inches(0.79)    # 2 cm = 0.79 inches
        section.bottom_margin = Inches(0.79) # 2 cm = 0.79 inches
        section.header_distance = Inches(0.49)  # 1.25 cm = 0.49 inches
        section.footer_distance = Inches(0.49)  # 1.25 cm = 0.49 inches
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Tiêu đề
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f"DANH SÁCH VỊ TRÍ CÁN BỘ NĂM {nam}")
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Đơn vị
        unit_para = doc.add_paragraph()
        unit_run = unit_para.add_run(f"Đơn vị: {don_vi}")
        unit_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 11 cột
        table = doc.add_table(rows=1, cols=11)
        table.style = None  # Bỏ style mặc định
        
        # Thiết lập borders cho toàn bộ table
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Thêm borders cho toàn bộ table
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Thiết lập độ rộng bảng
        table_width_inches = 24.0 - 0.4  # 24 inches - margins
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))  # 1440 twips per inch
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Độ rộng cột
        col_widths = [
            Inches(0.5),   # TT
            Inches(4.0),   # Họ và tên...
            Inches(1.2),   # Cấp Bậc
            Inches(1.8),   # Chức, đơn vị
            Inches(1.0),   # CM Quân
            Inches(1.2),   # Vào Đảng
            Inches(2.5),   # Chức vụ chiến đấu
            Inches(2.0),   # Qua trường
            Inches(0.8),   # VH SK
            Inches(1.0),   # DT TG
            Inches(7.0),   # Thông tin gia đình
        ]
        
        for idx, width in enumerate(col_widths):
            if idx < len(table.columns):
                table.columns[idx].width = width
        
        # Header row
        header_cells = table.rows[0].cells
        headers = [
            'TT',
            'Họ và tên\nSinh (tuổi)\nQuê quán - trú quán\nSHSQ',
            'Cấp Bậc\n(Tháng\nnăm\nnhận)',
            'Chức,\nđơn vị\n(Tháng\nnăm\nnhận)',
            'CM\nQuân\n(Tháng\nnăm)',
            'Vào Đảng:\nChính thức',
            'Chức vụ chiến\nđấu (Thời gian)\nChức vụ đã qua\n(Thời gian)',
            'Qua trường\n(Ngành, thời\ngian, kết quả)',
            'VH\nSK',
            'DT\nTG',
            'Thông tin gia đình:\nBố đẻ: sinh, nghề\nMẹ đẻ: sinh, nghề\nNơi ở hiện nay\nBố vợ: sinh, nghề\nMẹ vợ: sinh, nghề\nNơi ở hiện nay\nVợ: sinh, nghề\nCon: sinh, nghề\nNơi ở hiện nay\nSĐT gia đình'
        ]
        
        for i, header_text in enumerate(headers):
            if i < len(header_cells):
                cell = header_cells[i]
                cell.text = header_text
                
                # Format header
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho header
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm dữ liệu
        for idx, p in enumerate(personnel_list, 1):
            row = table.add_row()
            cells = row.cells
            
            # Tính tuổi
            try:
                if p.ngaySinh:
                    birth_year = int(p.ngaySinh.split('/')[-1])
                    age = 2025 - birth_year
                else:
                    age = ''
            except:
                age = ''
            
            # Cột 1: TT
            cells[0].text = f"{idx:02d}"
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cột 2: Họ và tên, Sinh (tuổi), Quê quán - trú quán, SHSQ
            # Format: "Triệu Văn Dũng\n19/11/1991 (34)\nThôn Tát Dài, xã Chợ Rã, tỉnh Thái Nguyên\nThôn Tam Trung, xã Tam Giang, tỉnh Đắk Lắk\n17022..."
            col2_parts = []
            if p.hoTen:
                col2_parts.append(p.hoTen)
            if p.ngaySinh:
                ngay_sinh_text = p.ngaySinh
                if age:
                    ngay_sinh_text += f" ({age})"  # Có khoảng trắng trước ngoặc
                col2_parts.append(ngay_sinh_text)
            if p.queQuan:
                col2_parts.append(p.queQuan)
            if p.truQuan:
                col2_parts.append(p.truQuan)
            if p.ghiChu:
                col2_parts.append(p.ghiChu)
            col2_text = "\n".join(col2_parts)
            cells[1].text = col2_text
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 3: Cấp Bậc (Tháng năm nhận) - Format: "4/5/2024" hoặc "H2/ 06/12/2024"
            col3_text = ""
            if p.capBac and p.ngayNhanCapBac:
                # Format: cấp bậc/ MM/YYYY (rút ngắn)
                ngay_nhan = p.ngayNhanCapBac
                if '/' in ngay_nhan:
                    parts = ngay_nhan.split('/')
                    if len(parts) == 3:
                        # DD/MM/YYYY -> MM/YYYY
                        thang_nam = f"{parts[1]}/{parts[2]}"
                        col3_text = f"{p.capBac}/{thang_nam}"
                    else:
                        col3_text = f"{p.capBac}/{ngay_nhan}"
                else:
                    col3_text = f"{p.capBac}/{ngay_nhan}"
            elif p.capBac:
                col3_text = p.capBac
            cells[2].text = col3_text
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 4: Chức, đơn vị (Tháng năm nhận) - Format: "CTV C3.d15, 5/2019"
            col4_text = ""
            if p.chucVu and p.donVi:
                chuc_don_vi = f"{p.chucVu} {p.donVi}"
                if p.ngayNhanChucVu:
                    # Lấy tháng/năm từ ngayNhanChucVu (DD/MM/YYYY -> MM/YYYY)
                    ngay_nhan = p.ngayNhanChucVu
                    if '/' in ngay_nhan:
                        parts = ngay_nhan.split('/')
                        if len(parts) == 3:
                            # DD/MM/YYYY -> MM/YYYY
                            thang_nam = f"{parts[1]}/{parts[2]}"
                            col4_text = f"{chuc_don_vi}, {thang_nam}"
                        elif len(parts) >= 2:
                            # MM/YYYY
                            thang_nam = f"{parts[0]}/{parts[1]}"
                            col4_text = f"{chuc_don_vi}, {thang_nam}"
                        else:
                            col4_text = f"{chuc_don_vi}, {ngay_nhan}"
                    else:
                        col4_text = f"{chuc_don_vi}, {ngay_nhan}"
                else:
                    col4_text = chuc_don_vi
            elif p.chucVu:
                col4_text = p.chucVu
            elif p.donVi:
                col4_text = p.donVi
            cells[3].text = col4_text
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 5: CM Quân (Tháng năm) - Format: "9/2012"
            cm_quan = p.cmQuan or p.nhapNgu or ''
            # Chuyển đổi format nếu cần (DD/MM/YYYY -> MM/YYYY)
            if cm_quan and '/' in cm_quan:
                parts = cm_quan.split('/')
                if len(parts) == 3:
                    # DD/MM/YYYY -> MM/YYYY
                    cm_quan = f"{parts[1]}/{parts[2]}"
                elif len(parts) >= 2:
                    # MM/YYYY hoặc YYYY
                    cm_quan = f"{parts[-2]}/{parts[-1]}"
            cells[4].text = cm_quan
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 6: Vào Đảng: Chính thức - Format: "5/2015\n5/2016"
            col6_parts = []
            if p.thongTinKhac.dang.ngayVao:
                # Chuyển DD/MM/YYYY -> MM/YYYY
                ngay_vao = p.thongTinKhac.dang.ngayVao
                if '/' in ngay_vao:
                    parts = ngay_vao.split('/')
                    if len(parts) == 3:
                        # DD/MM/YYYY -> MM/YYYY
                        ngay_vao = f"{parts[1]}/{parts[2]}"
                    elif len(parts) >= 2:
                        ngay_vao = f"{parts[-2]}/{parts[-1]}"
                col6_parts.append(ngay_vao)
            if p.thongTinKhac.dang.ngayChinhThuc:
                ngay_chinh_thuc = p.thongTinKhac.dang.ngayChinhThuc
                if '/' in ngay_chinh_thuc:
                    parts = ngay_chinh_thuc.split('/')
                    if len(parts) == 3:
                        # DD/MM/YYYY -> MM/YYYY
                        ngay_chinh_thuc = f"{parts[1]}/{parts[2]}"
                    elif len(parts) >= 2:
                        ngay_chinh_thuc = f"{parts[-2]}/{parts[-1]}"
                col6_parts.append(ngay_chinh_thuc)
            col6_text = "\n".join(col6_parts)
            cells[5].text = col6_text
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 7: Chức vụ chiến đấu (Thời gian) Chức vụ đã qua (Thời gian)
            # Format: "CTVP/c10.d9.e66.fl0 (8/2017-5/2019), - CTV/c3.d15.f10 (5/2019-)"
            col7_parts = []
            if p.chucVuChienDau:
                chien_dau_text = p.chucVuChienDau
                if p.thoiGianChucVuChienDau:
                    chien_dau_text += f" ({p.thoiGianChucVuChienDau})"
                col7_parts.append(chien_dau_text)
            elif p.chucVu and p.donVi:
                # Fallback: dùng chức vụ và đơn vị hiện tại
                col7_parts.append(f"{p.chucVu}/{p.donVi}")
            
            if p.chucVuDaQua:
                da_qua_text = f"-{p.chucVuDaQua}"
                if p.thoiGianChucVuDaQua:
                    da_qua_text += f" ({p.thoiGianChucVuDaQua})"
                col7_parts.append(da_qua_text)
            
            col7_text = ", ".join(col7_parts) if col7_parts else ''
            cells[6].text = col7_text
            cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 8: Qua trường (Ngành, thời gian, kết quả)
            # Format: "SQCT (Chính trị BCHT, 9/2012-7/2017-khá)"
            col8_text = p.quaTruong or ''
            if col8_text and (p.nganhHoc or p.thoiGianDaoTao or p.ketQuaDaoTao):
                col8_text += " ("
                parts = []
                if p.nganhHoc:
                    parts.append(p.nganhHoc)
                if p.thoiGianDaoTao:
                    parts.append(p.thoiGianDaoTao)
                if p.ketQuaDaoTao:
                    # Kết quả được nối với dấu "-" trước thời gian
                    if p.thoiGianDaoTao:
                        # Nếu có thời gian, kết quả nối sau dấu "-"
                        parts[-1] = f"{parts[-1]}-{p.ketQuaDaoTao}"
                    else:
                        parts.append(f"-{p.ketQuaDaoTao}")
                col8_text += ", ".join(parts)
                col8_text += ")"
            cells[7].text = col8_text
            cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 9: VH SK
            cells[8].text = p.trinhDoVanHoa or ''
            cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 10: DT TG
            col10_text = f"{p.danToc or ''}\n{p.tonGiao or 'Không'}"
            cells[9].text = col10_text
            cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cột 11: Thông tin gia đình
            gia_dinh_info = []
            if db_service:
                try:
                    nguoi_than_list = db_service.get_nguoi_than_by_personnel(p.id)
                    
                    bo_de = []
                    me_de = []
                    bo_vo = []
                    me_vo = []
                    vo = []
                    con = []
                    dia_chi_bo_me = None
                    dia_chi_bo_me_vo = None
                    dia_chi_vo_con = None
                    sdt_gia_dinh = None
                    
                    for nguoi_than in nguoi_than_list:
                        moi_quan_he = (nguoi_than.moiQuanHe or '').lower()
                        ho_ten = nguoi_than.hoTen or ''
                        ngay_sinh = nguoi_than.ngaySinh or ''
                        dia_chi = nguoi_than.diaChi or ''
                        so_dt = nguoi_than.soDienThoai or ''
                        noi_dung = nguoi_than.noiDung or ''
                        
                        # Lấy năm sinh
                        nam_sinh = ""
                        if ngay_sinh:
                            try:
                                if '/' in ngay_sinh:
                                    nam_sinh = ngay_sinh.split('/')[-1]
                                else:
                                    nam_sinh = ngay_sinh[:4] if len(ngay_sinh) >= 4 else ngay_sinh
                            except:
                                nam_sinh = ""
                        
                        nghe = noi_dung if noi_dung else "LN"
                        
                        if ho_ten:
                            if nam_sinh:
                                info_str = f"{ho_ten}, {nam_sinh}, {nghe}"
                            else:
                                info_str = f"{ho_ten}, {nghe}"
                            
                            if 'bố' in moi_quan_he or 'cha' in moi_quan_he:
                                if 'vợ' in moi_quan_he or 'vo' in moi_quan_he:
                                    bo_vo.append(f"Bố vợ: - {info_str}")
                                else:
                                    bo_de.append(f"Bố đẻ: - {info_str}")
                            elif 'mẹ' in moi_quan_he or 'me' in moi_quan_he:
                                if 'vợ' in moi_quan_he or 'vo' in moi_quan_he:
                                    me_vo.append(f"Mẹ vợ: - {info_str}")
                                else:
                                    me_de.append(f"Mẹ đẻ: - {info_str}")
                            elif 'vợ' in moi_quan_he or 'vo' in moi_quan_he:
                                vo.append(f"Vợ: - {info_str}")
                            elif 'con' in moi_quan_he:
                                con.append(f"Con: - {info_str}")
                        
                        if dia_chi:
                            if ('bố' in moi_quan_he or 'cha' in moi_quan_he or 'mẹ' in moi_quan_he or 'me' in moi_quan_he) and not ('vợ' in moi_quan_he or 'vo' in moi_quan_he):
                                if not dia_chi_bo_me:
                                    dia_chi_bo_me = dia_chi
                            elif ('bố' in moi_quan_he or 'cha' in moi_quan_he or 'mẹ' in moi_quan_he or 'me' in moi_quan_he) and ('vợ' in moi_quan_he or 'vo' in moi_quan_he):
                                if not dia_chi_bo_me_vo:
                                    dia_chi_bo_me_vo = dia_chi
                            elif 'vợ' in moi_quan_he or 'vo' in moi_quan_he or 'con' in moi_quan_he:
                                if not dia_chi_vo_con:
                                    dia_chi_vo_con = dia_chi
                        
                        if so_dt and not sdt_gia_dinh:
                            sdt_gia_dinh = so_dt
                    
                    gia_dinh_info.extend(bo_de)
                    gia_dinh_info.extend(me_de)
                    if dia_chi_bo_me:
                        gia_dinh_info.append(f"Nơi ở hiện nay: {dia_chi_bo_me}")
                    elif p.queQuan:
                        gia_dinh_info.append(f"Nơi ở hiện nay: {p.queQuan}")
                    gia_dinh_info.extend(bo_vo)
                    gia_dinh_info.extend(me_vo)
                    if dia_chi_bo_me_vo:
                        gia_dinh_info.append(f"Nơi ở hiện nay: {dia_chi_bo_me_vo}")
                    gia_dinh_info.extend(vo)
                    gia_dinh_info.extend(con)
                    if dia_chi_vo_con:
                        gia_dinh_info.append(f"Nơi ở hiện nay: {dia_chi_vo_con}")
                    if sdt_gia_dinh:
                        gia_dinh_info.append(f"SĐT gia đình: {sdt_gia_dinh}")
                    elif p.soDienThoaiLienHe:
                        gia_dinh_info.append(f"SĐT gia đình: {p.soDienThoaiLienHe}")
                except:
                    pass
            
            cells[10].text = "\n".join(gia_dinh_info) if gia_dinh_info else ''
            cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Format tất cả cells
            for cell in cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho từng cell (đảm bảo có borders)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm phần chữ ký ở cuối
        doc.add_paragraph()  # Khoảng trống
        doc.add_paragraph()  # Khoảng trống
        
        # Phần chữ ký
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # "CHÍNH TRỊ VIÊN"
        chinh_tri_vien_label = signature_para.add_run("CHÍNH TRỊ VIÊN")
        chinh_tri_vien_label.bold = True
        chinh_tri_vien_label.font.size = Pt(11)
        chinh_tri_vien_label.font.name = 'Times New Roman'
        chinh_tri_vien_label._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Xuống dòng
        signature_para.add_run("\n")
        
        # Tên chính trị viên
        chinh_tri_vien_name = signature_para.add_run(chinh_tri_vien)
        chinh_tri_vien_name.bold = True
        chinh_tri_vien_name.font.size = Pt(11)
        chinh_tri_vien_name.font.name = 'Times New Roman'
        chinh_tri_vien_name._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("Cần cài đặt python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

