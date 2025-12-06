"""
Hàm xuất Word cho Ban Chấp Hành Chi Đoàn
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_ban_chap_hanh_chi_doan(personnel_list: List[Personnel],
                                        don_vi: str = "Đại đội 3",
                                        tieu_doan: str = "TIỂU ĐOÀN 38",
                                        dia_diem: str = "Đăk Lăk",
                                        ngay_thang_nam: str = "",
                                        ten_bi_thu: str = "",
                                        db_service=None) -> bytes:
    """
    Xuất danh sách Ban Chấp Hành Chi Đoàn ra file Word với 12 cột giống mẫu
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang Letter (27.94 cm x 21.59 cm)
        section = doc.sections[0]
        section.page_width = Inches(11.0)  # 27.94 cm = 11 inches
        section.page_height = Inches(8.5)  # 21.59 cm = 8.5 inches
        section.left_margin = Inches(0.79)  # 2 cm
        section.right_margin = Inches(0.79)  # 2 cm
        section.top_margin = Inches(0.79)  # 2 cm
        section.bottom_margin = Inches(0.79)  # 2 cm
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: ĐOÀN CƠ SỞ TIỂU ĐOÀN 38, CHI ĐOÀN ĐẠI ĐỘI 3
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_run1 = left_cell.paragraphs[0].add_run("ĐOÀN CƠ SỞ " + tieu_doan)
        left_run1.font.size = Pt(11)
        left_run1.font.name = 'Times New Roman'
        left_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_cell.paragraphs[0].add_run("\n")
        left_run2 = left_cell.paragraphs[0].add_run(f"CHI ĐOÀN {don_vi.upper()}")
        left_run2.font.size = Pt(11)
        left_run2.font.name = 'Times New Roman'
        left_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Cột phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM, Độc lập - Tự do - Hạnh phúc
        right_cell = header_table.rows[0].cells[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        right_run1 = right_cell.paragraphs[0].add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        right_run1.font.size = Pt(11)
        right_run1.font.name = 'Times New Roman'
        right_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        right_cell.paragraphs[0].add_run("\n")
        right_run2 = right_cell.paragraphs[0].add_run("Độc lập - Tự do - Hạnh phúc")
        right_run2.font.size = Pt(11)
        right_run2.font.name = 'Times New Roman'
        right_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run2.underline = True
        
        # Merge cells cho dòng 2 (date line)
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
        date_cell = header_table.rows[1].cells[0]
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_cell.paragraphs[0].clear()
        
        if not ngay_thang_nam:
            ngay_thang_nam = datetime.now().strftime("%d/%m/%Y")
        
        date_text = f"{dia_diem}, ngày tháng năm {datetime.now().year}"
        date_run = date_cell.paragraphs[0].add_run(date_text)
        date_run.font.size = Pt(11)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Bỏ borders cho header table
        for row in header_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tiêu đề: DANH SÁCH
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DANH SÁCH")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle: Ban chấp hành Chi đoàn Đại đội 3
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Ban chấp hành Chi đoàn {don_vi}")
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        subtitle_run.underline = True
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 12 cột
        table = doc.add_table(rows=1, cols=12)
        table.style = None  # Bỏ style mặc định
        
        # Thiết lập borders cho toàn bộ table
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Thêm borders cho toàn bộ table
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Độ rộng cột
        col_widths = [
            Inches(0.4),   # TT
            Inches(1.5),   # Họ và tên
            Inches(0.8),   # N.T.năm sinh
            Inches(0.6),   # Cấp bậc
            Inches(0.7),   # Chức vụ
            Inches(0.7),   # Nhập ngũ
            Inches(0.7),   # Đơn vị
            Inches(0.6),   # Văn hóa
            Inches(1.0),   # Dân tộc, Tôn giáo
            Inches(1.0),   # Đảng, Đoàn
            Inches(2.0),   # Quê quán Trú quán
            Inches(0.8),   # Ghi chú
        ]
        
        for idx, width in enumerate(col_widths):
            if idx < len(table.columns):
                table.columns[idx].width = width
        
        # Header row
        header_cells = table.rows[0].cells
        headers = [
            'TT',
            'Họ và tên',
            'N.T.năm sinh',
            'Cấp bậc',
            'Chức vụ',
            'Nhập ngũ',
            'Đơn vị',
            'Văn hóa',
            'Dân tộc,\nTôn giáo',
            'Đảng,\nĐoàn',
            'Quê quán\nTrú quán',
            'Ghi chú'
        ]
        
        for i, header_text in enumerate(headers):
            if i < len(header_cells):
                cell = header_cells[i]
                cell.text = header_text
                
                # Format header
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho header
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm dữ liệu
        for idx, p in enumerate(personnel_list, 1):
            row = table.add_row()
            cells = row.cells
            
            # Format tất cả cells
            for cell in cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho từng cell
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Cột 1: TT
            cells[0].text = str(idx)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cột 2: Họ và tên
            cells[1].text = p.hoTen or ''
            
            # Cột 3: N.T.năm sinh
            cells[2].text = p.ngaySinh or ''
            
            # Cột 4: Cấp bậc
            cells[3].text = p.capBac or ''
            
            # Cột 5: Chức vụ
            cells[4].text = p.chucVu or ''
            
            # Cột 6: Nhập ngũ
            nhap_ngu = p.nhapNgu or ''
            # Format: MM/YYYY nếu có
            if '/' in nhap_ngu and len(nhap_ngu.split('/')) == 3:
                parts = nhap_ngu.split('/')
                nhap_ngu = f"{parts[1]}/{parts[2]}"
            cells[5].text = nhap_ngu
            
            # Cột 7: Đơn vị
            cells[6].text = p.donVi or ''
            
            # Cột 8: Văn hóa
            cells[7].text = p.trinhDoVanHoa or ''
            
            # Cột 9: Dân tộc, Tôn giáo
            dan_toc = p.danToc or ''
            ton_giao = p.tonGiao or 'Không'
            cells[8].text = f"{dan_toc}\n{ton_giao}"
            
            # Cột 10: Đảng, Đoàn
            ngay_vao_dang = ''
            ngay_chinh_thuc_dang = ''
            if p.thongTinKhac and p.thongTinKhac.dang:
                ngay_vao_dang = p.thongTinKhac.dang.ngayVao or ''
                ngay_chinh_thuc_dang = p.thongTinKhac.dang.ngayChinhThuc or ''
                # Format: DD/MM/YYYY -> DD/MM/YYYY
                if ngay_vao_dang and '/' in ngay_vao_dang:
                    parts = ngay_vao_dang.split('/')
                    if len(parts) == 3:
                        ngay_vao_dang = f"{parts[0]}/{parts[1]}/{parts[2]}"
                if ngay_chinh_thuc_dang and '/' in ngay_chinh_thuc_dang:
                    parts = ngay_chinh_thuc_dang.split('/')
                    if len(parts) == 3:
                        ngay_chinh_thuc_dang = f"{parts[0]}/{parts[1]}/{parts[2]}"
            
            ngay_vao_doan = ''
            if p.thongTinKhac and p.thongTinKhac.doan:
                ngay_vao_doan = p.thongTinKhac.doan.ngayVao or ''
                if ngay_vao_doan and '/' in ngay_vao_doan:
                    parts = ngay_vao_doan.split('/')
                    if len(parts) == 3:
                        ngay_vao_doan = f"{parts[0]}/{parts[1]}/{parts[2]}"
            
            # Lấy chức vụ đoàn từ bảng ban_chap_hanh_chi_doan
            chuc_vu_doan = ''
            if db_service:
                chuc_vu_doan = db_service.get_chuc_vu_doan(p.id)
                if not chuc_vu_doan and p.thongTinKhac and p.thongTinKhac.doan:
                    chuc_vu_doan = p.thongTinKhac.doan.chucVuDoan or ''
            
            # Format: Đảng: ngay_vao_dang\nngay_chinh_thuc_dang\nĐoàn: ngay_vao_doan
            dang_doan_parts = []
            if ngay_vao_dang:
                dang_doan_parts.append(ngay_vao_dang)
            if ngay_chinh_thuc_dang:
                dang_doan_parts.append(ngay_chinh_thuc_dang)
            if ngay_vao_doan:
                dang_doan_parts.append(ngay_vao_doan)
            
            cells[9].text = "\n".join(dang_doan_parts) if dang_doan_parts else ''
            
            # Cột 11: Quê quán Trú quán
            que_quan = p.queQuan or ''
            tru_quan = p.truQuan or ''
            if que_quan and tru_quan:
                cells[10].text = f"{que_quan}\n{tru_quan}"
            elif que_quan:
                cells[10].text = que_quan
            elif tru_quan:
                cells[10].text = tru_quan
            else:
                cells[10].text = ''
            
            # Cột 12: Ghi chú
            # Lấy chức vụ đoàn làm ghi chú nếu có
            if chuc_vu_doan:
                cells[11].text = chuc_vu_doan
            else:
                cells[11].text = p.ghiChu or ''
        
        # Footer: T/M BCH CHI ĐOÀN, BÍ THƯ, tên
        doc.add_paragraph()  # Khoảng trống
        
        footer_table = doc.add_table(rows=3, cols=1)
        footer_table.autofit = False
        footer_cell = footer_table.rows[0].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_cell.paragraphs[0].clear()
        footer_run1 = footer_cell.paragraphs[0].add_run("T/M BCH CHI ĐOÀN")
        footer_run1.font.size = Pt(11)
        footer_run1.font.name = 'Times New Roman'
        footer_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        footer_cell = footer_table.rows[1].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_cell.paragraphs[0].clear()
        footer_run2 = footer_cell.paragraphs[0].add_run("BÍ THƯ")
        footer_run2.font.size = Pt(11)
        footer_run2.font.name = 'Times New Roman'
        footer_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        footer_cell = footer_table.rows[2].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_cell.paragraphs[0].clear()
        if ten_bi_thu:
            footer_run3 = footer_cell.paragraphs[0].add_run(ten_bi_thu)
            footer_run3.font.size = Pt(11)
            footer_run3.font.name = 'Times New Roman'
            footer_run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            footer_run3.underline = True
        
        # Bỏ borders cho footer table
        for row in footer_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("Cần cài đặt python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

