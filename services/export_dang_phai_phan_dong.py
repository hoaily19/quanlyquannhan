"""
Hàm xuất Word cho Quân nhân có người thân tham gia đảng phái phản động
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_dang_phai_phan_dong(personnel_list: List[Personnel],
                                      tieu_doan: str = "TIỂU ĐOÀN 38",
                                      dai_doi: str = "ĐẠI ĐỘI 3",
                                      dia_diem: str = "Đắk Lắk",
                                      chinh_tri_vien: str = "",
                                      nam: str = "2025",
                                      db_service=None) -> bytes:
    """
    Xuất danh sách Quân nhân có người thân tham gia đảng phái phản động ra file Word
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang A4 Landscape (29.7 cm x 21 cm)
        section = doc.sections[0]
        section.page_width = Inches(11.69)  # 29.7 cm
        section.page_height = Inches(8.27)  # 21 cm
        section.left_margin = Inches(0.906)  # 2.3 cm
        section.right_margin = Inches(0.787)  # 2 cm
        section.top_margin = Inches(0.433)  # 1.1 cm
        section.bottom_margin = Inches(0.689)  # 1.75 cm
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: TIỂU ĐOÀN 38, ĐẠI ĐỘI 3
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_run1 = left_cell.paragraphs[0].add_run(tieu_doan)
        left_run1.font.size = Pt(12)
        left_run1.font.name = 'Times New Roman'
        left_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_cell.paragraphs[0].add_run("\n")
        left_run2 = left_cell.paragraphs[0].add_run(dai_doi.upper())
        left_run2.font.size = Pt(12)
        left_run2.font.name = 'Times New Roman'
        left_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        left_run2.underline = True
        
        # Cột phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
        right_cell = header_table.rows[0].cells[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        right_run1 = right_cell.paragraphs[0].add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        right_run1.font.size = Pt(12)
        right_run1.font.name = 'Times New Roman'
        right_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        right_cell.paragraphs[0].add_run("\n")
        right_run2 = right_cell.paragraphs[0].add_run("Độc lập - Tự do – Hạnh phúc")
        right_run2.font.size = Pt(12)
        right_run2.font.name = 'Times New Roman'
        right_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run2.underline = True
        
        # Merge cells cho dòng 2 (date line)
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
        date_cell = header_table.rows[1].cells[0]
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_cell.paragraphs[0].clear()
        
        date_text = f"{dia_diem}, ngày tháng năm {nam}"
        date_run = date_cell.paragraphs[0].add_run(date_text)
        date_run.font.size = Pt(12)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Bỏ borders cho header table
        for row in header_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tiêu đề: DANH SÁCH
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DANH SÁCH")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle: Quân nhân có người thân tham gia đảng phái, tổ chức chính trị...
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_text = "Quân nhân có người thân tham gia đảng phái, tổ chức chính trị, hội nhóm phản động, khủng bố; mít tinh, biểu tình trái pháp luật"
        subtitle_run = subtitle.add_run(subtitle_text)
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        subtitle_run.underline = True
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 7 cột
        table = doc.add_table(rows=1, cols=7)
        table.style = None
        
        # Thiết lập borders
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table width
        table_width_inches = 11.69 - 0.906 - 0.787
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Header row
        header_row = table.rows[0]
        headers = [
            'STT',
            'Họ và tên\n(Ngày, tháng năm sinh)\nNhập ngũ\nCấp bậc-chức vụ',
            'Đơn vị\n(ghi rõ từ c,d,e,f)',
            'Quê quán\nChỗ ở hiện nay',
            'Họ và tên người thân\n(Năm sinh, nghề nghiệp)\nMối quan hệ với quân nhân',
            'Nội dung người thân tham gia\n(Diễn biến, thời gian)',
            'Xử lý của địa phương'
        ]
        
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Xử lý header có xuống dòng
            if '\n' in header_text:
                parts = header_text.split('\n')
                for i, part in enumerate(parts):
                    if i > 0:
                        cell.paragraphs[0].add_run('\n')
                    run = cell.paragraphs[0].add_run(part)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            else:
                run = cell.paragraphs[0].add_run(header_text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Set cell width
            if idx == 0:  # STT
                cell.width = Inches(0.4)
            elif idx == 1:  # Họ và tên
                cell.width = Inches(2.0)
            elif idx == 2:  # Đơn vị
                cell.width = Inches(0.8)
            elif idx == 3:  # Quê quán
                cell.width = Inches(1.5)
            elif idx == 4:  # Họ và tên người thân
                cell.width = Inches(2.0)
            elif idx == 5:  # Nội dung
                cell.width = Inches(2.0)
            else:  # Xử lý của địa phương
                cell.width = Inches(1.5)
        
        # Data rows
        for idx, person in enumerate(personnel_list, 1):
            row = table.add_row()
            
            # Lấy danh sách người thân
            nguoi_than_list = []
            if db_service:
                try:
                    nguoi_than_list = db_service.get_nguoi_than_by_personnel(person.id)
                except:
                    pass
            
            # Nếu không có người thân, vẫn tạo 1 dòng trống
            if not nguoi_than_list:
                nguoi_than_list = [None]
            
            # Tạo 1 dòng cho mỗi người thân (hoặc 1 dòng trống nếu không có)
            for nguoi_than in nguoi_than_list:
                if nguoi_than != nguoi_than_list[0]:  # Nếu không phải người thân đầu tiên, tạo dòng mới
                    row = table.add_row()
                
                # STT
                cell = row.cells[0]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell.paragraphs[0].add_run(str(idx))
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Họ và tên (với thông tin bổ sung)
                cell = row.cells[1]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                ho_ten_full = person.hoTen or ''
                if person.ngaySinh:
                    ho_ten_full += f"\n({person.ngaySinh})"
                if person.nhapNgu:
                    ho_ten_full += f"\n{person.nhapNgu}"
                cb_cv = f"{person.capBac or ''}-{person.chucVu or ''}".strip('-')
                if cb_cv:
                    ho_ten_full += f"\n{cb_cv}"
                run = cell.paragraphs[0].add_run(ho_ten_full)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Đơn vị
                cell = row.cells[2]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell.paragraphs[0].add_run(person.donVi or '')
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Quê quán / Chỗ ở hiện nay
                cell = row.cells[3]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                que_quan = person.queQuan or ''
                tru_quan = person.truQuan or ''
                que_tru = f"{que_quan}\n{tru_quan}".strip('\n')
                run = cell.paragraphs[0].add_run(que_tru)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Họ và tên người thân (với thông tin bổ sung)
                cell = row.cells[4]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                if nguoi_than:
                    ho_ten_nt = nguoi_than.hoTen or ''
                    # Lấy năm sinh từ ngày sinh
                    nam_sinh = ""
                    if nguoi_than.ngaySinh:
                        try:
                            if '/' in nguoi_than.ngaySinh:
                                parts = nguoi_than.ngaySinh.split('/')
                                nam_sinh = parts[-1] if len(parts) >= 3 else (parts[1] if len(parts) == 2 else parts[0][:4])
                            else:
                                nam_sinh = nguoi_than.ngaySinh[:4] if len(nguoi_than.ngaySinh) >= 4 else nguoi_than.ngaySinh
                        except:
                            nam_sinh = ""
                    
                    moi_quan_he = nguoi_than.moiQuanHe or ''
                    
                    nguoi_than_full = ho_ten_nt
                    if nam_sinh:
                        nguoi_than_full += f"\n({nam_sinh})"
                    if moi_quan_he:
                        nguoi_than_full += f"\n{moi_quan_he}"
                else:
                    nguoi_than_full = ""
                run = cell.paragraphs[0].add_run(nguoi_than_full)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Nội dung người thân tham gia
                cell = row.cells[5]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                noi_dung = ""
                if nguoi_than:
                    noi_dung = nguoi_than.noiDung or ''
                run = cell.paragraphs[0].add_run(noi_dung)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Xử lý của địa phương (để trống)
                cell = row.cells[6]
                cell.paragraphs[0].clear()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = cell.paragraphs[0].add_run("")
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        doc.add_paragraph()  # Khoảng trống
        
        # Footer: CHÍNH TRỊ VIÊN
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("CHÍNH TRỊ VIÊN")
        footer_run.bold = True
        footer_run.font.size = Pt(11)
        footer_run.font.name = 'Times New Roman'
        footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        if chinh_tri_vien:
            footer2 = doc.add_paragraph()
            footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer2_run = footer2.add_run(chinh_tri_vien)
            footer2_run.font.size = Pt(11)
            footer2_run.font.name = 'Times New Roman'
            footer2_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

