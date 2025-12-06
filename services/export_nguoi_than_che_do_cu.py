"""
Hàm xuất Word cho Quân nhân có người thân tham gia chế độ cũ
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_nguoi_than_che_do_cu(personnel_list: List[Personnel],
                                      tieu_doan: str = "TIỂU ĐOÀN 38",
                                      dai_doi: str = "ĐẠI ĐỘI 3",
                                      dia_diem: str = "Đắk Lắk",
                                      chinh_tri_vien: str = "",
                                      db_service=None) -> bytes:
    """
    Xuất danh sách Quân nhân có người thân tham gia chế độ cũ ra file Word
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang A4 Landscape (29.7 cm x 21 cm)
        section = doc.sections[0]
        section.page_width = Inches(11.69)  # 29.7 cm
        section.page_height = Inches(8.27)  # 21 cm
        section.left_margin = Inches(0.906)  # 2.3 cm
        section.right_margin = Inches(0.787)  # 2 cm
        section.top_margin = Inches(0.433)  # 1.1 cm
        section.bottom_margin = Inches(0.689)  # 1.75 cm
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: TIỂU ĐOÀN 38, ĐẠI ĐỘI 3
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_run1 = left_cell.paragraphs[0].add_run(tieu_doan)
        left_run1.font.size = Pt(12)
        left_run1.font.name = 'Times New Roman'
        left_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_cell.paragraphs[0].add_run("\n")
        left_run2 = left_cell.paragraphs[0].add_run(dai_doi.upper())
        left_run2.font.size = Pt(12)
        left_run2.font.name = 'Times New Roman'
        left_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        left_run2.underline = True
        
        # Cột phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
        right_cell = header_table.rows[0].cells[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        right_run1 = right_cell.paragraphs[0].add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        right_run1.font.size = Pt(12)
        right_run1.font.name = 'Times New Roman'
        right_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        right_cell.paragraphs[0].add_run("\n")
        right_run2 = right_cell.paragraphs[0].add_run("Độc lập - Tự do – Hạnh phúc")
        right_run2.font.size = Pt(12)
        right_run2.font.name = 'Times New Roman'
        right_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run2.underline = True
        
        # Merge cells cho dòng 2 (date line)
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
        date_cell = header_table.rows[1].cells[0]
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_cell.paragraphs[0].clear()
        
        date_text = f"{dia_diem}, ngày tháng năm {datetime.now().strftime('%Y')}"
        date_run = date_cell.paragraphs[0].add_run(date_text)
        date_run.font.size = Pt(12)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Bỏ borders cho header table
        for row in header_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tiêu đề: DANH SÁCH
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DANH SÁCH")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle: Quân nhân có người thân tham gia chế độ cũ
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Quân nhân có người thân tham gia chế độ cũ")
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        subtitle_run.underline = True
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 7 cột chính (cột THAM GIA sẽ có 5 sub-columns)
        # Tổng cộng: STT(1) + Họ và tên(1) + Đơn vị(1) + THAM GIA(5) + Họ tên người thân(1) + Quan hệ(1) + Đã cải tạo(1) = 11 cột
        table = doc.add_table(rows=2, cols=11)  # 2 hàng header
        table.style = None
        
        # Thiết lập borders
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table width
        table_width_inches = 11.69 - 0.906 - 0.787
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Độ rộng cột
        col_widths = [
            Inches(0.4),   # STT
            Inches(2.0),   # Họ và tên
            Inches(0.6),   # Đơn vị
            Inches(0.5),   # Ngụy quân
            Inches(0.5),   # Ngụy quyền
            Inches(0.8),   # Nợ máu/không nợ máu
            Inches(1.2),   # Quê quán
            Inches(1.2),   # Chỗ ở hiện nay
            Inches(2.0),   # Họ tên người thân
            Inches(0.8),   # Quan hệ
            Inches(1.0),   # Đã cải tạo/chưa cải tạo
        ]
        
        for idx, width in enumerate(col_widths):
            if idx < len(table.columns):
                table.columns[idx].width = width
        
        # Header row 1 - merge các cột
        header_row1 = table.rows[0]
        # STT
        header_row1.cells[0].text = 'STT'
        # Họ và tên (merge 1 cột)
        header_row1.cells[1].text = 'Họ và tên\n(Ngày, tháng, năm sinh)\nNhập ngũ\nCấp bậc-chức vụ'
        # Đơn vị
        header_row1.cells[2].text = 'ĐƠN VỊ\n(ghi rõ từ c,d,e,f)'
        # THAM GIA - merge 5 cột
        tham_gia_cell = header_row1.cells[3]
        for i in range(4, 8):
            tham_gia_cell.merge(header_row1.cells[i])
        tham_gia_cell.text = 'THAM GIA'
        # Họ tên người thân
        header_row1.cells[8].text = 'Họ tên người thân\n(năm sinh, nghề nghiệp)\nNội dung cụ thể về hoạt động tham gia CĐC'
        # Quan hệ
        header_row1.cells[9].text = 'Quan hệ với quân nhân'
        # Đã cải tạo
        header_row1.cells[10].text = 'Đã cải tạo/chưa cải tạo'
        
        # Header row 2 - sub-headers cho THAM GIA
        header_row2 = table.rows[1]
        # STT - merge với row 1
        header_row1.cells[0].merge(header_row2.cells[0])
        # Họ và tên - merge với row 1
        header_row1.cells[1].merge(header_row2.cells[1])
        # Đơn vị - merge với row 1
        header_row1.cells[2].merge(header_row2.cells[2])
        # THAM GIA sub-columns
        header_row2.cells[3].text = 'Ngụy quân'
        header_row2.cells[4].text = 'Ngụy quyền'
        header_row2.cells[5].text = 'Nợ máu/không nợ máu'
        header_row2.cells[6].text = 'Quê quán'
        header_row2.cells[7].text = 'Chỗ ở hiện nay'
        # Họ tên người thân - merge với row 1
        header_row1.cells[8].merge(header_row2.cells[8])
        # Quan hệ - merge với row 1
        header_row1.cells[9].merge(header_row2.cells[9])
        # Đã cải tạo - merge với row 1
        header_row1.cells[10].merge(header_row2.cells[10])
        
        # Format headers
        for row_idx in range(2):
            for cell in table.rows[row_idx].cells:
                # Format header - nền trắng, chữ đen
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Vertical alignment center
                tcPr = cell._element.tcPr
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm dữ liệu
        for idx, person in enumerate(personnel_list, 1):
            row = table.add_row()
            cells = row.cells
            
            # STT
            cells[0].text = str(idx)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Họ và tên với thông tin bổ sung
            ho_ten_text = person.hoTen or ''
            if person.ngaySinh:
                ho_ten_text += f"\n({person.ngaySinh})"
            if person.nhapNgu:
                ho_ten_text += f"\nNhập ngũ: {person.nhapNgu}"
            cb_cv = f"{person.capBac or ''}-{person.chucVu or ''}".strip('-')
            if cb_cv:
                ho_ten_text += f"\n{cb_cv}"
            cells[1].text = ho_ten_text
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Đơn vị
            cells[2].text = person.donVi or ''
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # THAM GIA
            cells[3].text = 'X' if person.thamGiaNguyQuan else ''
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[4].text = 'X' if person.thamGiaNguyQuyen else ''
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[5].text = person.thamGiaNoMau or ''
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[6].text = person.queQuan or ''
            cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[7].text = person.truQuan or ''
            cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Họ tên người thân
            nguoi_than_info = ""
            quan_he = ""
            if db_service and person.id:
                try:
                    nguoi_than_list = db_service.get_nguoi_than_by_personnel(person.id)
                    if nguoi_than_list:
                        nt = nguoi_than_list[0]
                        ho_ten_nt = nt.hoTen or ''
                        ngay_sinh_nt = nt.ngaySinh or ''
                        noi_dung_nt = nt.noiDung or ''
                        
                        # Lấy năm sinh
                        nam_sinh = ""
                        if ngay_sinh_nt:
                            try:
                                if '/' in ngay_sinh_nt:
                                    parts = ngay_sinh_nt.split('/')
                                    nam_sinh = parts[-1] if len(parts) >= 3 else (parts[1] if len(parts) == 2 else parts[0][:4])
                                else:
                                    nam_sinh = ngay_sinh_nt[:4] if len(ngay_sinh_nt) >= 4 else ngay_sinh_nt
                            except:
                                nam_sinh = ""
                        
                        if nam_sinh:
                            nguoi_than_info = f"{ho_ten_nt} ({nam_sinh}, {noi_dung_nt})"
                        else:
                            nguoi_than_info = f"{ho_ten_nt} ({noi_dung_nt})"
                        
                        quan_he = nt.moiQuanHe or ''
                except:
                    pass
            
            cells[8].text = nguoi_than_info
            cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Quan hệ
            cells[9].text = quan_he
            cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Đã cải tạo
            cells[10].text = person.daCaiTao or ''
            cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Format cells
            for cell in cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm dòng tổng
        total_row = table.add_row()
        total_cells = total_row.cells
        total_cells[0].text = "Tổng"
        total_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[1].text = str(len(personnel_list))
        total_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format tổng
        for cell in total_cells:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'FFFFFF')
            shading_elm.set(qn('w:val'), 'clear')
            cell._element.get_or_add_tcPr().append(shading_elm)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        doc.add_paragraph()  # Khoảng trống
        
        # Chữ ký (căn phải)
        if chinh_tri_vien:
            signature = doc.add_paragraph()
            signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            signature_run = signature.add_run("CHÍNH TRỊ VIÊN\n")
            signature_run.bold = True
            signature_run.font.size = Pt(12)
            signature_run.font.name = 'Times New Roman'
            signature_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            name_run = signature.add_run(chinh_tri_vien)
            name_run.bold = True
            name_run.font.size = Pt(12)
            name_run.font.name = 'Times New Roman'
            name_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("Cần cài đặt python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")


