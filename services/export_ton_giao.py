"""
Hàm xuất Word cho Quân Nhân Theo Tôn Giáo
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_ton_giao(personnel_list: List[Personnel],
                          don_vi: str = "Đại đội 3",
                          tieu_doan: str = "TIỂU ĐOÀN 38",
                          dia_diem: str = "Đắk Lắk",
                          ngay_thang_nam: str = "",
                          chinh_tri_vien: str = "") -> bytes:
    """
    Xuất danh sách Quân Nhân Theo Tôn Giáo ra file Word với format giống mẫu
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang Letter (27.94 cm x 21.59 cm)
        section = doc.sections[0]
        section.page_width = Inches(11.0)  # 27.94 cm = 11 inches
        section.page_height = Inches(8.5)  # 21.59 cm = 8.5 inches
        section.left_margin = Inches(0.79)  # 2 cm
        section.right_margin = Inches(0.79)  # 2 cm
        section.top_margin = Inches(0.79)  # 2 cm
        section.bottom_margin = Inches(0.79)  # 2 cm
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: TIỂU ĐOÀN 38, ĐẠI ĐỘI 3
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_run1 = left_cell.paragraphs[0].add_run(tieu_doan)
        left_run1.font.size = Pt(11)
        left_run1.font.name = 'Times New Roman'
        left_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_cell.paragraphs[0].add_run("\n")
        left_run2 = left_cell.paragraphs[0].add_run(don_vi.upper())
        left_run2.font.size = Pt(11)
        left_run2.font.name = 'Times New Roman'
        left_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Cột phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM, Độc lập - Tự do - Hạnh phúc
        right_cell = header_table.rows[0].cells[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        right_run1 = right_cell.paragraphs[0].add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        right_run1.font.size = Pt(11)
        right_run1.font.name = 'Times New Roman'
        right_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run1.bold = True
        
        right_cell.paragraphs[0].add_run("\n")
        right_run2 = right_cell.paragraphs[0].add_run("Độc lập - Tự do - Hạnh phúc")
        right_run2.font.size = Pt(11)
        right_run2.font.name = 'Times New Roman'
        right_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        right_run2.italic = True
        right_run2.underline = True
        
        # Merge cells cho dòng 2 (date line)
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
        date_cell = header_table.rows[1].cells[0]
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_cell.paragraphs[0].clear()
        
        if not ngay_thang_nam:
            ngay_thang_nam = datetime.now().strftime("%d/%m/%Y")
        
        date_text = f"{dia_diem}, ngày tháng năm {datetime.now().year}"
        date_run = date_cell.paragraphs[0].add_run(date_text)
        date_run.font.size = Pt(11)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Bỏ borders cho header table
        for row in header_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tiêu đề: DANH SÁCH
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DANH SÁCH")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle: Quân nhân theo tôn giáo
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Quân nhân theo tôn giáo")
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        subtitle_run.bold = True
        subtitle_run.underline = True
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 7 cột chính (cột Tôn giáo sẽ có 4 sub-columns)
        # Thực tế là 10 cột: STT, Họ tên, N.ngũ, CB-CV, Đ.vị, Quê quán, Chỗ ở, Thiên chúa giáo, Phật giáo, Tin lành, Công giáo, Ghi chú
        # Nhưng để đơn giản, tôi sẽ tạo 12 cột: STT, Họ tên, N.ngũ, CB, CV, Đ.vị, Quê quán, Chỗ ở, Thiên chúa giáo, Phật giáo, Tin lành, Công giáo, Ghi chú
        # Thực ra theo mẫu: STT, Họ tên (Ngày, tháng năm sinh), N.ngũ CB - CV, Đ. vị, Quê quán Chỗ ở hiện nay, Tôn giáo (4 sub), Ghi chú
        # Tổng cộng: 1 + 1 + 1 + 1 + 1 + 4 + 1 = 10 cột
        table = doc.add_table(rows=1, cols=10)
        table.style = None  # Bỏ style mặc định
        
        # Thiết lập borders cho toàn bộ table
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Thêm borders cho toàn bộ table
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Độ rộng cột
        col_widths = [
            Inches(0.4),   # STT
            Inches(2.0),   # Họ tên (Ngày, tháng năm sinh)
            Inches(1.0),   # N.ngũ CB - CV
            Inches(0.6),   # Đ. vị
            Inches(2.5),   # Quê quán Chỗ ở hiện nay
            Inches(0.8),   # Thiên chúa giáo
            Inches(0.8),   # Phật giáo
            Inches(0.8),   # Tin lành
            Inches(0.8),   # Công giáo
            Inches(0.8),   # Ghi chú
        ]
        
        for idx, width in enumerate(col_widths):
            if idx < len(table.columns):
                table.columns[idx].width = width
        
        # Header row
        header_cells = table.rows[0].cells
        headers = [
            'STT',
            'Họ tên\n(Ngày, tháng\nnăm sinh)',
            'N.ngũ\nCB - CV',
            'Đ. vị',
            'Quê quán\nChỗ ở hiện nay',
            'Thiên chúa\ngiáo',
            'Phật giáo',
            'Tin lành',
            'Công giáo',
            'Ghi chú'
        ]
        
        for i, header_text in enumerate(headers):
            if i < len(header_cells):
                cell = header_cells[i]
                cell.text = header_text
                
                # Format header
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho header
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Đếm số lượng theo từng tôn giáo
        thien_chua_count = 0
        phat_giao_count = 0
        tin_lanh_count = 0
        cong_giao_count = 0
        
        # Thêm dữ liệu
        for idx, p in enumerate(personnel_list, 1):
            row = table.add_row()
            cells = row.cells
            
            # Format tất cả cells
            for cell in cells:
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho từng cell
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Cột 1: STT
            cells[0].text = str(idx)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cột 2: Họ tên (Ngày, tháng năm sinh)
            ho_ten = p.hoTen or ''
            ngay_sinh = p.ngaySinh or ''
            if ho_ten and ngay_sinh:
                cells[1].text = f"{ho_ten}\n({ngay_sinh})"
            elif ho_ten:
                cells[1].text = ho_ten
            else:
                cells[1].text = ''
            
            # Cột 3: N.ngũ CB - CV
            nhap_ngu = p.nhapNgu or ''
            # Format: MM/YYYY nếu có
            if '/' in nhap_ngu and len(nhap_ngu.split('/')) == 3:
                parts = nhap_ngu.split('/')
                nhap_ngu = f"{parts[1]}/{parts[2]}"
            
            cap_bac = p.capBac or ''
            chuc_vu = p.chucVu or ''
            if nhap_ngu and cap_bac and chuc_vu:
                cells[2].text = f"{nhap_ngu}\n{cap_bac}-{chuc_vu}"
            elif nhap_ngu and cap_bac:
                cells[2].text = f"{nhap_ngu}\n{cap_bac}"
            elif cap_bac and chuc_vu:
                cells[2].text = f"{cap_bac}-{chuc_vu}"
            else:
                cells[2].text = nhap_ngu or cap_bac or chuc_vu or ''
            
            # Cột 4: Đ. vị
            cells[3].text = p.donVi or ''
            
            # Cột 5: Quê quán Chỗ ở hiện nay
            que_quan = p.queQuan or ''
            tru_quan = p.truQuan or ''
            if que_quan and tru_quan:
                cells[4].text = f"{que_quan}\n{tru_quan}"
            elif que_quan:
                cells[4].text = que_quan
            elif tru_quan:
                cells[4].text = tru_quan
            else:
                cells[4].text = ''
            
            # Cột 6-9: Tôn giáo (Thiên chúa giáo, Phật giáo, Tin lành, Công giáo)
            ton_giao = (p.tonGiao or '').strip().lower()
            cells[5].text = ''  # Thiên chúa giáo
            cells[6].text = ''  # Phật giáo
            cells[7].text = ''  # Tin lành
            cells[8].text = ''  # Công giáo
            
            if 'thiên chúa' in ton_giao or 'công giáo' in ton_giao:
                if 'công giáo' in ton_giao:
                    cells[8].text = 'X'  # Công giáo
                    cong_giao_count += 1
                else:
                    cells[5].text = 'X'  # Thiên chúa giáo
                    thien_chua_count += 1
            elif 'phật' in ton_giao:
                cells[6].text = 'X'  # Phật giáo
                phat_giao_count += 1
            elif 'tin lành' in ton_giao or 'tinlanh' in ton_giao:
                cells[7].text = 'X'  # Tin lành
                tin_lanh_count += 1
            
            # Căn giữa cho các cột tôn giáo
            for i in range(5, 9):
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Cột 10: Ghi chú
            cells[9].text = p.ghiChu or ''
        
        # Thêm dòng tổng
        total_row = table.add_row()
        total_cells = total_row.cells
        
        # Format tổng row
        for cell in total_cells:
            tcPr = cell._element.get_or_add_tcPr()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'FFFFFF')
            shading_elm.set(qn('w:val'), 'clear')
            tcPr.append(shading_elm)
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcPr.append(border)
            
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Điền dữ liệu tổng
        total_cells[0].text = 'Tổng'
        total_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[1].text = str(len(personnel_list))
        total_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[2].text = ''
        total_cells[3].text = ''
        total_cells[4].text = ''
        total_cells[5].text = str(thien_chua_count)
        total_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[6].text = str(phat_giao_count)
        total_cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[7].text = str(tin_lanh_count)
        total_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[8].text = str(cong_giao_count)
        total_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[9].text = ''
        
        # Footer: CHÍNH TRỊ VIÊN, tên
        doc.add_paragraph()  # Khoảng trống
        
        footer_table = doc.add_table(rows=2, cols=1)
        footer_table.autofit = False
        footer_cell = footer_table.rows[0].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_cell.paragraphs[0].clear()
        footer_run1 = footer_cell.paragraphs[0].add_run("CHÍNH TRỊ VIÊN")
        footer_run1.font.size = Pt(11)
        footer_run1.font.name = 'Times New Roman'
        footer_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        footer_run1.bold = True
        
        footer_cell = footer_table.rows[1].cells[0]
        footer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_cell.paragraphs[0].clear()
        if chinh_tri_vien:
            footer_run2 = footer_cell.paragraphs[0].add_run(chinh_tri_vien)
            footer_run2.font.size = Pt(11)
            footer_run2.font.name = 'Times New Roman'
            footer_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            footer_run2.bold = True
        
        # Bỏ borders cho footer table
        for row in footer_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("Cần cài đặt python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")


