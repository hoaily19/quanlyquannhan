"""
Hàm xuất Word cho Tổ công tác dân vận
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_to_dan_van(personnel_list: List[Personnel],
                           don_vi: str = "Đại đội 3",
                           tieu_doan: str = "TIỂU ĐOÀN 38",
                           dia_diem: str = "Đắk Lắk",
                           ngay_thang_nam: str = "",
                           chinh_tri_vien: str = "",
                           db_service=None) -> bytes:
    """
    Xuất danh sách Tổ công tác dân vận ra file Word với format giống mẫu
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header: TIỂU ĐOÀN 38 và ĐẠI ĐỘI 3
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header_para.add_run(f"{tieu_doan}\n{don_vi.upper()}")
        header_run.font.size = Pt(12)
        header_run.font.name = 'Times New Roman'
        header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header phải: CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
        # Tạo paragraph với tab để căn phải
        header_right_para = doc.add_paragraph()
        header_right_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_right_run = header_right_para.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc")
        header_right_run.font.size = Pt(12)
        header_right_run.font.name = 'Times New Roman'
        header_right_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Ngày tháng năm
        if not ngay_thang_nam:
            ngay_thang_nam = datetime.now().strftime("%d/%m/%Y")
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_para.add_run(f"{dia_diem}, ngày {ngay_thang_nam.split('/')[0] if '/' in ngay_thang_nam else ''} tháng {ngay_thang_nam.split('/')[1] if '/' in ngay_thang_nam and len(ngay_thang_nam.split('/')) > 1 else ''} năm {ngay_thang_nam.split('/')[2] if '/' in ngay_thang_nam and len(ngay_thang_nam.split('/')) > 2 else datetime.now().strftime('%Y')}")
        date_run.font.size = Pt(12)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Tiêu đề: DANH SÁCH TỔ CÔNG TÁC DÂN VẬN
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("DANH SÁCH TỔ CÔNG TÁC DÂN VẬN")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 9 cột
        table = doc.add_table(rows=1, cols=9)
        table.style = None
        
        # Thiết lập borders cho toàn bộ table
        tbl = table._element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)
        
        # Header row
        header_row = table.rows[0]
        header_cells = header_row.cells
        
        headers = ['TT', 'Họ và tên', 'Cấp bậc\nChức vụ', 'Đơn vị', 
                  'Dân tộc\nTôn giáo', 'Trình độ\nvăn hóa', 'Ngoại ngữ', 
                  'Tiếng DTTS', 'Ghi chú']
        
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            # Format header
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Set cell width
            if i == 0:  # TT
                cell.width = Inches(0.4)
            elif i == 1:  # Họ và tên
                cell.width = Inches(1.2)
            elif i == 2:  # Cấp bậc/Chức vụ
                cell.width = Inches(0.8)
            elif i == 3:  # Đơn vị
                cell.width = Inches(0.5)
            elif i == 4:  # Dân tộc/Tôn giáo
                cell.width = Inches(1.0)
            elif i == 5:  # Trình độ văn hóa
                cell.width = Inches(0.7)
            elif i == 6:  # Ngoại ngữ
                cell.width = Inches(0.7)
            elif i == 7:  # Tiếng DTTS
                cell.width = Inches(0.7)
            else:  # Ghi chú
                cell.width = Inches(1.0)
        
        # Thêm dữ liệu
        for idx, person in enumerate(personnel_list, 1):
            row = table.add_row()
            cells = row.cells
            
            # TT
            cells[0].text = str(idx)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Họ và tên
            cells[1].text = person.hoTen or ''
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Cấp bậc/Chức vụ
            cb_cv = f"{person.capBac or ''}/{person.chucVu or ''}".strip('/')
            cells[2].text = cb_cv
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Đơn vị
            cells[3].text = person.donVi or ''
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dân tộc/Tôn giáo
            dan_toc = person.danToc or ''
            ton_giao = person.tonGiao or 'Không'
            cells[4].text = f"{dan_toc} / {ton_giao}"
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Trình độ văn hóa
            cells[5].text = person.trinhDoVanHoa or ''
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ngoại ngữ
            cells[6].text = person.ngoaiNgu or ''
            cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Tiếng DTTS
            cells[7].text = person.tiengDTTS or ''
            cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ghi chú - lấy từ ghi chú riêng của tab
            ghi_chu = ''
            if db_service:
                ghi_chu = db_service.get_to_dan_van_ghi_chu(person.id)
            cells[8].text = ghi_chu
            cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Format font cho tất cả cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Footer: CHÍNH TRỊ VIÊN
        doc.add_paragraph()  # Khoảng trống
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_para.add_run("CHÍNH TRỊ VIÊN")
        footer_run.font.size = Pt(12)
        footer_run.font.name = 'Times New Roman'
        footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        if chinh_tri_vien:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            name_run = name_para.add_run(chinh_tri_vien)
            name_run.font.size = Pt(12)
            name_run.font.name = 'Times New Roman'
            name_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Lưu vào bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
        
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

