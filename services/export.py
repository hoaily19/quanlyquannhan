"""
Service xuất dữ liệu ra CSV và PDF
"""

import csv
import io
import sys
from typing import List
from datetime import datetime
from pathlib import Path

# Thêm thư mục gốc vào path
sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


class ExportService:
    """Service xuất file"""
    
    @staticmethod
    def to_csv(personnel_list: List[Personnel]) -> str:
        """
        Chuyển đổi danh sách quân nhân sang CSV
        Returns:
            Chuỗi CSV
        """
        if not personnel_list:
            return ""
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Header - Đầy đủ các trường với chính tả đúng
        headers = [
            'Họ và Tên', 'Họ Tên Thường Dùng', 'Ngày Sinh', 'Cấp Bậc', 
            'Ngày Nhận Cấp Bậc', 'Chức Vụ', 'Ngày Nhận Chức Vụ', 'Đơn Vị',
            'Nhập Ngũ', 'Xuất Ngũ', 'Quê Quán', 'Nơi Cư Trú', 'Dân Tộc', 'Tôn Giáo',
            'Trình Độ Văn Hóa', 'Thành Phần Gia Đình', 'Qua Trường', 'Ngành Học',
            'Cấp Học', 'Thời Gian Đào Tạo', 'Liên Hệ Khi Cần', 'Số Điện Thoại Liên Hệ',
            'Họ Tên Cha', 'Họ Tên Mẹ', 'Họ Tên Vợ', 'Ghi Chú',
            'Ngày Vào Đảng', 'Ngày Chính Thức Đảng', 'Chức Vụ Đảng',
            'Ngày Vào Đoàn', 'Chức Vụ Đoàn',
            'Chế Độ Cũ', 'Yếu Tố Nước Ngoài'
        ]
        writer.writerow(headers)
        
        # Dữ liệu - Đầy đủ các trường
        for person in personnel_list:
            row = [
                person.hoTen or '',
                person.hoTenThuongDung or '',
                person.ngaySinh or '',
                person.capBac or '',
                person.ngayNhanCapBac or '',
                person.chucVu or '',
                person.ngayNhanChucVu or '',
                person.donVi or '',
                person.nhapNgu or '',
                person.xuatNgu or '',
                person.queQuan or '',
                person.truQuan or '',
                person.danToc or '',
                person.tonGiao or '',
                person.trinhDoVanHoa or '',
                person.thanhPhanGiaDinh or '',
                person.quaTruong or '',
                person.nganhHoc or '',
                person.capHoc or '',
                person.thoiGianDaoTao or '',
                person.lienHeKhiCan or '',
                person.soDienThoaiLienHe or '',
                person.hoTenCha or '',
                person.hoTenMe or '',
                person.hoTenVo or '',
                person.ghiChu or '',
                person.thongTinKhac.dang.ngayVao or '',
                person.thongTinKhac.dang.ngayChinhThuc or '',
                person.thongTinKhac.dang.chucVuDang or '',
                person.thongTinKhac.doan.ngayVao or '',
                person.thongTinKhac.doan.chucVuDoan or '',
                'Có' if person.thongTinKhac.cdCu else 'Không',
                'Có' if person.thongTinKhac.yeuToNN else 'Không',
            ]
            writer.writerow(row)
        
        return output.getvalue()
    
    @staticmethod
    def to_pdf(personnel_list: List[Personnel], title: str = "Danh Sách Quân Nhân") -> bytes:
        """
        Chuyển đổi danh sách quân nhân sang PDF
        Returns:
            Bytes của file PDF
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Tiêu đề
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=colors.HexColor('#4CAF50'),
                spaceAfter=30,
                alignment=1  # Center
            )
            story.append(Paragraph(title, title_style))
            story.append(Paragraph(f"Tổng số: {len(personnel_list)} quân nhân", styles['Normal']))
            story.append(Spacer(1, 0.5*cm))
            
            # Bảng dữ liệu
            if personnel_list:
                data = [['STT', 'Họ và Tên', 'Cấp Bậc', 'Chức Vụ', 'Đơn Vị', 'Dân Tộc']]
                
                for idx, person in enumerate(personnel_list, 1):
                    data.append([
                        str(idx),
                        person.hoTen or '',
                        person.capBac or '',
                        person.chucVu or '',
                        person.donVi or '',
                        person.danToc or '',
                    ])
                
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4CAF50')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                ]))
                
                story.append(table)
            
            # Footer
            story.append(Spacer(1, 1*cm))
            story.append(Paragraph(
                f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
                styles['Normal']
            ))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            # Fallback nếu không có reportlab
            raise ImportError("Cần cài đặt reportlab: pip install reportlab")

    @staticmethod
    def filter_ethnic_minority(personnel_list: List[Personnel]) -> List[Personnel]:
        """
        Lọc quân nhân là người đồng bào dân tộc thiểu số
        (Loại trừ dân tộc Kinh)
        """
        # Danh sách dân tộc thiểu số phổ biến
        ethnic_minorities = [
            'Nùng', 'Ede', 'Mông', 'Ja Rai', 'Sán chay', 'Thái', 'Tày', 'Ba na',
            'Ê đê', 'Gia Rai', 'Mường', 'Khmer', 'Hoa', 'Dao', 'Chăm', 'Sán Dìu',
            'H\'Mông', 'Hmông', 'Gia Rai', 'Jarai', 'Bana', 'Ba Na'
        ]
        
        # Chuẩn hóa danh sách (lowercase để so sánh không phân biệt hoa thường)
        ethnic_minorities_lower = [e.lower().strip() for e in ethnic_minorities]
        
        filtered = []
        for person in personnel_list:
            dan_toc = (person.danToc or '').strip()
            if dan_toc:
                # Nếu không phải Kinh và có trong danh sách dân tộc thiểu số
                if dan_toc.lower() not in ['kinh', 'việt']:
                    # Kiểm tra xem có khớp với danh sách không
                    if any(minority in dan_toc.lower() or dan_toc.lower() in minority 
                           for minority in ethnic_minorities_lower):
                        filtered.append(person)
                    # Hoặc nếu không phải Kinh thì cũng tính là dân tộc thiểu số
                    elif dan_toc.lower() not in ['kinh', 'việt', 'việt nam']:
                        filtered.append(person)
        
        return filtered
    
    @staticmethod
    def to_word_docx(personnel_list: List[Personnel], 
                     tieu_doan: str = "TIỂU ĐOÀN 38",
                     dai_doi: str = "ĐẠI ĐỘI 3",
                     dia_diem: str = "Đắk Lắk",
                     chinh_tri_vien: str = "") -> bytes:
        """
        Xuất danh sách quân nhân ra file Word (.docx) với định dạng như ảnh
        Tự động lấy danh sách dân tộc từ database
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            doc = Document()
            
            # Thiết lập độ rộng trang và margins theo cấu hình Page Setup từ ảnh
            # Paper: Width 29.7 cm (11.69 inches), Height 21 cm (8.27 inches) - A4 Landscape
            # Margins: Top 1.1 cm, Bottom 1.75 cm, Left 2.3 cm, Right 2 cm
            # Header/Footer: 1.27 cm (0.5 inches) from edge
            section = doc.sections[0]
            section.page_width = Inches(11.69)  # 29.7 cm = 11.69 inches (ngang)
            section.page_height = Inches(8.27)  # 21 cm = 8.27 inches (dọc)
            section.left_margin = Inches(0.906)  # 2.3 cm = 0.906 inches
            section.right_margin = Inches(0.787)  # 2 cm = 0.787 inches
            section.top_margin = Inches(0.433)  # 1.1 cm = 0.433 inches
            section.bottom_margin = Inches(0.689)  # 1.75 cm = 0.689 inches
            section.header_distance = Inches(0.5)  # 1.27 cm = 0.5 inches
            section.footer_distance = Inches(0.5)  # 1.27 cm = 0.5 inches
            
            # Thiết lập font mặc định
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Header với 2 cột (trái và phải)
            header_table = doc.add_table(rows=2, cols=2)
            header_table.autofit = False
            
            # Cột trái: Tiểu đoàn và Đại đội
            left_cell = header_table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_para.add_run(tieu_doan).font.size = Pt(12)
            left_para = left_cell.add_paragraph()
            dai_doi_run = left_para.add_run(dai_doi)
            dai_doi_run.font.size = Pt(12)
            dai_doi_run.underline = True
            
            # Cột phải: Cộng hòa XHCN Việt Nam
            right_cell = header_table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").font.size = Pt(12)
            right_para = right_cell.add_paragraph()
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc_lap_run = right_para.add_run("Độc lập - Tự do – Hạnh phúc")
            doc_lap_run.font.size = Pt(12)
            doc_lap_run.underline = True
            
            # Merge hàng header
            header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
            
            # Tiêu đề chính "DANH SÁCH" (giữa)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("DANH SÁCH")
            title_run.bold = True
            title_run.font.size = Pt(16)
            
            # Phụ đề "Quân nhân là người đồng bào dân tộc thiểu số" (giữa, gạch chân)
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run("Quân nhân là người đồng bào dân tộc thiểu số")
            subtitle_run.bold = True
            subtitle_run.font.size = Pt(12)
            subtitle_run.underline = True
            
            # Ngày tháng (căn phải)
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"{dia_diem}, ngày {datetime.now().strftime('%d')} tháng {datetime.now().strftime('%m')} năm {datetime.now().strftime('%Y')}")
            date_run.font.size = Pt(12)
            
            doc.add_paragraph()  # Khoảng trống
            
            # Lấy danh sách dân tộc từ dữ liệu thực tế (loại trừ Kinh)
            ethnic_set = set()
            for person in personnel_list:
                dan_toc = (person.danToc or '').strip()
                if dan_toc and dan_toc.lower() not in ['kinh', 'việt', 'việt nam']:
                    ethnic_set.add(dan_toc)
            
            # Sắp xếp danh sách dân tộc
            ethnic_groups = sorted(list(ethnic_set))
            
            # Đếm số lượng theo từng dân tộc
            ethnic_counts = {group: 0 for group in ethnic_groups}
            
            # Tính số cột: TT(1) + HỌ TÊN(1) + CB/CV(1) + ĐƠN VỊ(1) + Dân tộc(n) + Quê quán/Trú quán(1) + GHI CHÚ(1)
            num_cols = 4 + len(ethnic_groups) + 2  # 4 cột cố định + số dân tộc + 2 cột cuối
            
            # Tạo bảng (không dùng style có màu, sẽ format thủ công)
            table = doc.add_table(rows=1, cols=num_cols)
            # Tắt autofit để kiểm soát độ rộng chính xác
            table.autofit = False
            
            # Thêm border đen cho toàn bộ bảng và thiết lập độ rộng bảng
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Thêm borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Màu đen
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Tính table width = page width - left margin - right margin
            # 11.69 - 0.906 - 0.787 = 9.997 inches ≈ 10 inches
            table_width_inches = 11.69 - 0.906 - 0.787
            
            # Set table width - thêm vào tblPr đã có
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))  # Convert to twips
            tblW.set(qn('w:type'), 'dxa')  # dxa = twentieths of a point
            tblPr.append(tblW)
            
            # Thiết lập độ rộng cột - Tổng độ rộng = table_width_inches
            # Điều chỉnh tỷ lệ để phù hợp với page width mới (11.69 inches thay vì 24 inches)
            col_widths = [
                Inches(0.4),   # TT
                Inches(1.5),   # HỌ TÊN
                Inches(0.6),   # CB, CV
                Inches(0.6),   # ĐƠN VỊ
            ]
            # Thêm width cho các cột dân tộc (mỗi cột nhỏ)
            for _ in ethnic_groups:
                col_widths.append(Inches(0.4))  # Mỗi cột dân tộc
            # Quê quán / Trú quán - tính toán để tổng = table_width_inches
            # Tính tổng các cột đã có
            total_so_far = 0.4 + 1.5 + 0.6 + 0.6 + (0.4 * len(ethnic_groups)) + 0.8
            que_quan_width = table_width_inches - total_so_far
            if que_quan_width < 1.0:  # Đảm bảo tối thiểu 1 inch
                que_quan_width = 1.0
            col_widths.append(Inches(que_quan_width))  # Quê quán / Trú quán (tự động tính)
            col_widths.append(Inches(0.8))  # GHI CHÚ
            
            # Set width cho các cột - đảm bảo tổng = 24 inches
            for idx, width in enumerate(col_widths):
                if idx < len(table.columns):
                    table.columns[idx].width = width
            
            # Đảm bảo bảng không tự động mở rộng - kiểm tra lại table width
            # tblPr và tblW đã được set ở trên, nhưng cần đảm bảo không bị ghi đè
            tbl = table._element
            existing_tblPr = tbl.tblPr
            if existing_tblPr is not None:
                # Kiểm tra xem tblW đã có chưa
                existing_tblW = existing_tblPr.find(qn('w:tblW'))
                if existing_tblW is None:
                    # Nếu chưa có, thêm vào
                    tblW = OxmlElement('w:tblW')
                    tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))
                    tblW.set(qn('w:type'), 'dxa')
                    existing_tblPr.append(tblW)
                else:
                    # Nếu đã có, cập nhật lại
                    existing_tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))
                    existing_tblW.set(qn('w:type'), 'dxa')
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['TT', 'HỌ TÊN (Ngày, tháng, năm sinh)', 'CB, CV', 'ĐƠN VỊ'] + ethnic_groups + ['Quê quán / Trú quán', 'GHI CHÚ']
            
            # Đảm bảo số cột header khớp với số cột bảng
            if len(headers) != num_cols:
                # Nếu thiếu, thêm cột trống
                while len(headers) < num_cols:
                    headers.append('')
                # Nếu thừa, cắt bớt
                headers = headers[:num_cols]
            
            for i, header_text in enumerate(headers):
                if i < len(header_cells):
                    cell = header_cells[i]
                    
                    # Format header - nền trắng, chữ đen
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFFFF')  # Trắng
                    shading_elm.set(qn('w:val'), 'clear')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Xác định loại header để format phù hợp
                    is_ethnic_col = i >= 4 and i < 4 + len(ethnic_groups)
                    is_que_quan_col = i == 4 + len(ethnic_groups)
                    is_ghi_chu_col = i == 4 + len(ethnic_groups) + 1
                    
                    # Xử lý text header - chia thành nhiều dòng nếu cần
                    if 'HỌ TÊN' in header_text:
                        # Header có thêm (Ngày, tháng, năm sinh)
                        cell.text = 'HỌ TÊN\n(Ngày, tháng, năm sinh)'
                    elif header_text == 'CB, CV':
                        cell.text = 'CB,\nCV'
                    elif header_text == 'ĐƠN VỊ':
                        cell.text = 'ĐƠN VỊ'  # Giữ nguyên, không chia dòng
                    elif header_text == 'Quê quán / Trú quán':
                        cell.text = 'Quê quán\nTrú quán'  # 2 dòng, không có dấu /
                    elif header_text == 'GHI CHÚ':
                        cell.text = 'GHI CHÚ'  # Giữ nguyên, không chia dòng
                    elif is_ethnic_col:
                        # Các cột dân tộc - giữ nguyên tên đầy đủ, sẽ xoay dọc
                        cell.text = header_text
                    else:
                        cell.text = header_text
                    
                    # Format header text
                    tcPr = cell._element.tcPr
                    
                    # Vertical alignment center cho tất cả header
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)
                    
                    # Xoay text dọc CHỈ cho các cột dân tộc
                    if is_ethnic_col:
                        # Thêm text direction (vertical) - xoay 90 độ
                        textDirection = OxmlElement('w:textDirection')
                        textDirection.set(qn('w:val'), 'tbRl')  # Top to bottom, right to left (vertical)
                        tcPr.append(textDirection)
                    # Các cột khác không có textDirection = giữ nguyên ngang
                    
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Màu đen
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Thêm dữ liệu
            for idx, person in enumerate(personnel_list, 1):
                row = table.add_row()
                cells = row.cells
                
                # TT
                cells[0].text = str(idx)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # HỌ TÊN (ngày sinh ở dòng dưới)
                ho_ten = person.hoTen or ''
                if person.ngaySinh:
                    cells[1].text = f"{ho_ten}\n{person.ngaySinh}"
                else:
                    cells[1].text = ho_ten
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # CB, CV
                cb_cv = f"{person.capBac or ''}/{person.chucVu or ''}".strip('/')
                cells[2].text = cb_cv
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # ĐƠN VỊ
                cells[3].text = person.donVi or ''
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Dân tộc - đánh dấu X vào cột tương ứng
                dan_toc = (person.danToc or '').strip()
                for i, ethnic in enumerate(ethnic_groups):
                    # So sánh chính xác hoặc chứa
                    if dan_toc == ethnic or dan_toc.lower() == ethnic.lower():
                        cells[4 + i].text = 'X'
                        cells[4 + i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        ethnic_counts[ethnic] += 1
                
                # Quê quán / Trú quán (cột sau các cột dân tộc) - format với dấu ; như mẫu
                que_quan = person.queQuan or ''
                tru_quan = person.truQuan or ''
                if que_quan and tru_quan:
                    dia_chi = f"{que_quan}; {tru_quan}"
                elif que_quan:
                    dia_chi = que_quan
                elif tru_quan:
                    dia_chi = tru_quan
                else:
                    dia_chi = ''
                que_quan_col = 4 + len(ethnic_groups)
                cells[que_quan_col].text = dia_chi
                cells[que_quan_col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # GHI CHÚ (cột cuối)
                ghi_chu_col = 4 + len(ethnic_groups) + 1
                cells[ghi_chu_col].text = person.ghiChu or ''
                cells[ghi_chu_col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Format font size và màu cho các cell (nền trắng, chữ đen)
                for cell in cells:
                    # Set background white
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFFFF')  # Trắng
                    shading_elm.set(qn('w:val'), 'clear')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Màu đen
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Thêm dòng tổng kết
            total_row = table.add_row()
            total_cells = total_row.cells
            
            # Format nền trắng cho tất cả cell tổng kết
            for cell in total_cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')  # Trắng
                shading_elm.set(qn('w:val'), 'clear')
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Cột "Tổng"
            if len(total_cells) > 0:
                total_cells[0].text = "Tổng"
                total_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if total_cells[0].paragraphs[0].runs:
                    total_cells[0].paragraphs[0].runs[0].bold = True
                    total_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                else:
                    run = total_cells[0].paragraphs[0].add_run("Tổng")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Tổng số quân nhân (cột HỌ TÊN)
            if len(total_cells) > 1:
                total_cells[1].text = str(len(personnel_list))
                total_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if total_cells[1].paragraphs[0].runs:
                    total_cells[1].paragraphs[0].runs[0].bold = True
                    total_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                else:
                    run = total_cells[1].paragraphs[0].add_run(str(len(personnel_list)))
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Tổng theo từng dân tộc
            for i, ethnic in enumerate(ethnic_groups):
                col_idx = 4 + i
                if col_idx < len(total_cells):
                    total_cells[col_idx].text = str(ethnic_counts[ethnic])
                    total_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if total_cells[col_idx].paragraphs[0].runs:
                        total_cells[col_idx].paragraphs[0].runs[0].bold = True
                        total_cells[col_idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        run = total_cells[col_idx].paragraphs[0].add_run(str(ethnic_counts[ethnic]))
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
            
            doc.add_paragraph()  # Khoảng trống
            
            # Chữ ký (căn phải như ảnh)
            if chinh_tri_vien:
                signature = doc.add_paragraph()
                signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                signature_run = signature.add_run("CHÍNH TRỊ VIÊN\n")
                signature_run.bold = True
                signature_run.font.size = Pt(12)
                name_run = signature.add_run(chinh_tri_vien)
                name_run.bold = True
                name_run.font.size = Pt(12)
            
            # Lưu vào buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            raise ImportError("Cần cài đặt python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"Lỗi khi xuất file Word: {str(e)}")
    
    @staticmethod
    def to_word_docx_by_units(unit, child_units_data: List[Dict],
                             tieu_doan: str = "TIỂU ĐOÀN 38",
                             dai_doi: str = "",
                             dia_diem: str = "Đắk Lắk",
                             chinh_tri_vien: str = "") -> bytes:
        """
        Xuất danh sách quân nhân theo đơn vị và tổ ra file Word (.docx)
        Định dạng: 5 cột (TT, Họ và tên, Cấp bậc, Chức vụ, Ghi chú)
        Có nhóm đại đội/trung đội, có tổ, mỗi tổ 3 người
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from models.unit import Unit
            
            doc = Document()
            
            # Thiết lập độ rộng trang và margins theo cấu hình Page Setup
            section = doc.sections[0]
            section.page_width = Inches(11.69)  # 29.7 cm = 11.69 inches (ngang)
            section.page_height = Inches(8.27)  # 21 cm = 8.27 inches (dọc)
            section.left_margin = Inches(0.906)  # 2.3 cm
            section.right_margin = Inches(0.787)  # 2 cm
            section.top_margin = Inches(0.433)  # 1.1 cm
            section.bottom_margin = Inches(0.689)  # 1.75 cm
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
            
            # Thiết lập font mặc định
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Header với 2 cột (trái và phải)
            header_table = doc.add_table(rows=2, cols=2)
            header_table.autofit = False
            
            # Cột trái: Tiểu đoàn và Đại đội
            left_cell = header_table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_para.add_run(tieu_doan).font.size = Pt(12)
            left_para = left_cell.add_paragraph()
            dai_doi_text = dai_doi if dai_doi else unit.ten
            dai_doi_run = left_para.add_run(dai_doi_text)
            dai_doi_run.font.size = Pt(12)
            dai_doi_run.underline = True
            
            # Cột phải: Cộng hòa XHCN Việt Nam
            right_cell = header_table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM").font.size = Pt(12)
            right_para = right_cell.add_paragraph()
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            doc_lap_run = right_para.add_run("Độc lập - Tự do – Hạnh phúc")
            doc_lap_run.font.size = Pt(12)
            doc_lap_run.underline = True
            
            # Merge hàng header
            header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
            
            # Tiêu đề chính "DANH SÁCH TỔ 3 NGƯỜI" (giữa)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("DANH SÁCH TỔ 3 NGƯỜI")
            title_run.bold = True
            title_run.font.size = Pt(16)
            
            # Ngày tháng (căn phải)
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"{dia_diem}, ngày {datetime.now().strftime('%d')} tháng {datetime.now().strftime('%m')} năm {datetime.now().strftime('%Y')}")
            date_run.font.size = Pt(12)
            
            doc.add_paragraph()  # Khoảng trống
            
            # Tạo bảng với 5 cột: TT, Họ và tên, Cấp bậc, Chức vụ, Ghi chú
            table = doc.add_table(rows=1, cols=5)
            table.autofit = False
            
            # Thiết lập độ rộng bảng
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Thêm borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Set table width
            table_width_inches = 11.69 - 0.906 - 0.787
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
            # Thiết lập độ rộng cột
            col_widths = [
                Inches(0.5),   # TT
                Inches(3.0),   # Họ và tên
                Inches(1.0),   # Cấp bậc
                Inches(1.5),   # Chức vụ
                Inches(2.0),   # Ghi chú
            ]
            for idx, width in enumerate(col_widths):
                if idx < len(table.columns):
                    table.columns[idx].width = width
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['TT', 'Họ và tên', 'Cấp bậc', 'Chức vụ', 'Ghi chú']
            
            for i, header_text in enumerate(headers):
                if i < len(header_cells):
                    cell = header_cells[i]
                    cell.text = header_text
                    
                    # Format header - nền trắng, chữ đen
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFFFF')
                    shading_elm.set(qn('w:val'), 'clear')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Vertical alignment center
                    tcPr = cell._element.tcPr
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)
                    
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Thêm hàng nhóm (đại đội/trung đội) - chỉ một lần
            group_name = unit.ten.upper()
            
            # Hàng nhóm (merge tất cả cột)
            group_row = table.add_row()
            group_cell = group_row.cells[0]
            # Merge tất cả cột
            for i in range(1, 5):
                group_cell.merge(group_row.cells[i])
            
            group_cell.text = group_name
            group_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Format nhóm
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'FFFFFF')
            shading_elm.set(qn('w:val'), 'clear')
            group_cell._element.get_or_add_tcPr().append(shading_elm)
            
            for run in group_cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Thêm dữ liệu theo tổ
            stt = 1
            
            for child_data in child_units_data:
                child_unit = child_data['to']
                personnel_list = child_data['personnel']
                
                # Hàng tổ (merge tất cả cột)
                to_row = table.add_row()
                to_cell = to_row.cells[0]
                for i in range(1, 5):
                    to_cell.merge(to_row.cells[i])
                
                to_cell.text = child_unit.ten
                to_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Format tổ
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                to_cell._element.get_or_add_tcPr().append(shading_elm)
                
                for run in to_cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Thêm quân nhân trong tổ (tối đa 3 người)
                for person in personnel_list[:3]:  # Chỉ lấy 3 người đầu
                    row = table.add_row()
                    cells = row.cells
                    
                    # TT
                    cells[0].text = str(stt)
                    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Họ và tên
                    cells[1].text = person.hoTen or ''
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cấp bậc
                    cells[2].text = person.capBac or ''
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Chức vụ
                    cells[3].text = person.chucVu or ''
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Ghi chú (lấy từ ghiChu hoặc để trống)
                    cells[4].text = person.ghiChu or ''
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Format font
                    for cell in cells:
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'FFFFFF')
                        shading_elm.set(qn('w:val'), 'clear')
                        cell._element.get_or_add_tcPr().append(shading_elm)
                        
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    
                    stt += 1
            
            doc.add_paragraph()  # Khoảng trống
            
            # Chữ ký (căn phải)
            if chinh_tri_vien:
                signature = doc.add_paragraph()
                signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                signature_run = signature.add_run("CHÍNH TRỊ VIÊN\n")
                signature_run.bold = True
                signature_run.font.size = Pt(12)
                name_run = signature.add_run(chinh_tri_vien)
                name_run.bold = True
                name_run.font.size = Pt(12)
            
            # Lưu vào buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            raise ImportError("Cần cài đặt python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"Lỗi khi xuất file Word: {str(e)}")