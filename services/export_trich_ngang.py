"""
Hàm xuất Word cho Trích ngang Đại đội
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_trich_ngang(personnel_list: List[Personnel],
                             tieu_doan: str = "TIỂU ĐOÀN 38",
                             dai_doi: str = "ĐẠI ĐỘI 3",
                             dia_diem: str = "Đăk Lăk",
                             nam: str = "2025",
                             db_service=None,
                             units_data: List[dict] = None) -> bytes:
    """
    Xuất danh sách Trích ngang Đại đội ra file Word với format theo ảnh
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang: A3 Landscape theo thông số ảnh
        # A3: 41.99 cm x 29.7 cm (Landscape)
        # Mirror margins: Top 2 cm, Bottom 2 cm, Inside 3 cm, Outside 1.5 cm
        # Header/Footer: 1 cm from edge
        section = doc.sections[0]
        section.page_width = Cm(41.99)  # A3 width
        section.page_height = Cm(29.7)  # A3 height
        section.left_margin = Cm(1.5)   # Outside margin
        section.right_margin = Cm(3.0)  # Inside margin (mirror)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: Tiểu đoàn và Đại đội
        left_cell = header_table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_run = left_para.add_run(tieu_doan)
        left_run.font.size = Pt(12)
        left_run.font.name = 'Times New Roman'
        left_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_para = left_cell.add_paragraph()
        dai_doi_run = left_para.add_run(dai_doi)
        dai_doi_run.font.size = Pt(12)
        dai_doi_run.font.name = 'Times New Roman'
        dai_doi_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Cột phải: Cộng hòa XHCN Việt Nam
        right_cell = header_table.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        right_run.font.size = Pt(12)
        right_run.font.name = 'Times New Roman'
        right_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        right_para = right_cell.add_paragraph()
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc_lap_run = right_para.add_run("Độc lập – Tự do – Hạnh phúc")
        doc_lap_run.font.size = Pt(12)
        doc_lap_run.font.name = 'Times New Roman'
        doc_lap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Hàng 2: Ngày tháng năm
        date_cell = header_table.rows[1].cells[0]
        date_cell.merge(header_table.rows[1].cells[1])
        date_para = date_cell.paragraphs[0]
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"{dia_diem}, ngày   tháng   năm {nam}")
        date_run.font.size = Pt(12)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Tiêu đề "DANH SÁCH Trích ngang Đại đội"
        doc.add_paragraph()  # Khoảng trống
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("DANH SÁCH Trích ngang Đại đội")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 13 cột (theo mẫu)
        table = doc.add_table(rows=1, cols=13)
        table.style = None  # Bỏ style mặc định
        
        # Thiết lập borders cho toàn bộ table
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Thêm borders cho toàn bộ table
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Độ rộng cột (tính theo tỷ lệ, tổng ~41.99 cm - margins = ~37.49 cm)
        col_widths = [
            Cm(1.0),   # 1. SỐ TT
            Cm(3.0),   # 2. Họ và tên khai sinh / Họ và tên thường dùng
            Cm(2.5),   # 3. Ngày tháng năm sinh / Cấp bậc / Ngày nhận
            Cm(1.8),   # 4. Chức vụ / Ngày nhận
            Cm(1.8),   # 5. Nhập ngũ / Xuất ngũ
            Cm(2.2),   # 6. N. vào đoàn / N.vào đảng Chính thức
            Cm(2.2),   # 7. Thành phần GĐ / Dân tộc / Tôn giáo
            Cm(1.2),   # 8. Văn hóa
            Cm(2.5),   # 9. Qua trường / Ngành học / Cấp học / Thời gian
            Cm(3.5),   # 10. Quê quán / Trú quán / Khi cần báo tin cho ai SĐT
            Cm(2.5),   # 11. Họ tên cha / Họ tên mẹ / Họ tên vợ
            Cm(1.8),   # 12. Đơn vị đang làm nhiệm vụ
            Cm(1.5),   # 13. Ghi chú
        ]
        
        for idx, width in enumerate(col_widths):
            if idx < len(table.columns):
                table.columns[idx].width = width
        
        # Header row - 13 cột theo mẫu
        header_cells = table.rows[0].cells
        headers = [
            'SỐ TT',
            'Họ và tên khai sinh\nHọ và tên thường dùng',
            'Ngày tháng năm sinh\nCấp bậc\nNgày nhận',
            'Chức vụ\nNgày nhận',
            'Nhập ngũ\nXuất ngũ',
            'N. vào đoàn\nN.vào đảng\nChính thức',
            'Thành phần GĐ\nDân tộc\nTôn giáo',
            'Văn hóa',
            'Qua trường\nNgành học\nCấp học\nThời gian',
            'Quê quán\nTrú quán\nKhi cần báo tin cho ai SĐT',
            'Họ tên cha\nHọ tên mẹ\nHọ tên vợ',
            'Đơn vị đang làm nhiệm vụ',
            'Ghi chú'
        ]
        
        for i, header_text in enumerate(headers):
            if i < len(header_cells):
                cell = header_cells[i]
                cell.text = header_text
                
                # Format header
                tcPr = cell._element.get_or_add_tcPr()
                
                # Shading - nền trắng
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFFFFF')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders cho header
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                # Vertical alignment center
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm dữ liệu - nếu có units_data thì nhóm theo đơn vị
        if units_data:
            # Nhóm theo đơn vị với sub-header
            stt_counter = 1
            for unit_group in units_data:
                unit = unit_group.get('to')
                personnel_in_unit = unit_group.get('personnel', [])
                
                if not personnel_in_unit:
                    continue
                
                # Thêm sub-header row cho đơn vị (merge tất cả 13 cột)
                header_row = table.add_row()
                header_cells = header_row.cells
                
                # Merge tất cả cells trong hàng header
                for i in range(1, 13):
                    header_cells[0].merge(header_cells[i])
                
                # Set text cho sub-header
                unit_name = unit.ten if unit else "Đơn vị"
                header_cells[0].text = unit_name
                header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Format sub-header
                tcPr = header_cells[0]._element.get_or_add_tcPr()
                
                # Shading - nền xám nhạt
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E0E0E0')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
                
                # Borders
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcPr.append(border)
                
                for paragraph in header_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Thêm dữ liệu quân nhân trong đơn vị này
                for p in personnel_in_unit:
                    row = table.add_row()
                    cells = row.cells
                    
                    # Format tất cả cells
                    for cell in cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        
                        # Shading - nền trắng
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), 'FFFFFF')
                        shading_elm.set(qn('w:val'), 'clear')
                        tcPr.append(shading_elm)
                        
                        # Borders cho từng cell
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:space'), '0')
                            border.set(qn('w:color'), '000000')
                            tcPr.append(border)
                        
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    
                    # Cột 1: STT
                    cells[0].text = str(stt_counter)
                    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    stt_counter += 1
                    
                    # Cột 2: Họ và tên khai sinh / Họ và tên thường dùng
                    ho_ten = p.hoTen or ''
                    if p.hoTenThuongDung:
                        ho_ten += f"\n{p.hoTenThuongDung}"
                    cells[1].text = ho_ten
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 3: Ngày tháng năm sinh / Cấp bậc / Ngày nhận
                    col3_parts = []
                    if p.ngaySinh:
                        col3_parts.append(p.ngaySinh)
                    if p.capBac:
                        col3_parts.append(p.capBac)
                    if p.ngayNhanCapBac:
                        # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                        ngay_nhan = p.ngayNhanCapBac
                        if '/' in ngay_nhan:
                            parts = ngay_nhan.split('/')
                            if len(parts) == 3:
                                ngay_nhan = f"{parts[1]}/{parts[2]}"
                        col3_parts.append(ngay_nhan)
                    cells[2].text = "\n".join(col3_parts)
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 4: Chức vụ / Ngày nhận
                    col4_parts = []
                    if p.chucVu:
                        col4_parts.append(p.chucVu)
                    if p.ngayNhanChucVu:
                        # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                        ngay_nhan = p.ngayNhanChucVu
                        if '/' in ngay_nhan:
                            parts = ngay_nhan.split('/')
                            if len(parts) == 3:
                                ngay_nhan = f"{parts[1]}/{parts[2]}"
                        col4_parts.append(ngay_nhan)
                    cells[3].text = "\n".join(col4_parts)
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 5: Nhập ngũ / Xuất ngũ
                    col5_parts = []
                    if p.nhapNgu:
                        # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                        nhap_ngu = p.nhapNgu
                        if '/' in nhap_ngu:
                            parts = nhap_ngu.split('/')
                            if len(parts) == 3:
                                nhap_ngu = f"{parts[1]}/{parts[2]}"
                        col5_parts.append(nhap_ngu)
                    if p.xuatNgu:
                        # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                        xuat_ngu = p.xuatNgu
                        if '/' in xuat_ngu:
                            parts = xuat_ngu.split('/')
                            if len(parts) == 3:
                                xuat_ngu = f"{parts[1]}/{parts[2]}"
                        col5_parts.append(xuat_ngu)
                    cells[4].text = "\n".join(col5_parts)
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 6: N. vào đoàn / N.vào đảng Chính thức
                    col6_parts = []
                    if p.thongTinKhac.doan.ngayVao:
                        # Giữ nguyên format DD/MM/YYYY
                        col6_parts.append(p.thongTinKhac.doan.ngayVao)
                    if p.thongTinKhac.dang.ngayVao:
                        # Giữ nguyên format DD/MM/YYYY
                        col6_parts.append(p.thongTinKhac.dang.ngayVao)
                    if p.thongTinKhac.dang.ngayChinhThuc:
                        # Giữ nguyên format DD/MM/YYYY
                        col6_parts.append(p.thongTinKhac.dang.ngayChinhThuc)
                    cells[5].text = "\n".join(col6_parts)
                    cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 7: Thành phần GĐ / Dân tộc / Tôn giáo
                    col7_parts = []
                    if p.thanhPhanGiaDinh:
                        col7_parts.append(p.thanhPhanGiaDinh)
                    if p.danToc:
                        col7_parts.append(p.danToc)
                    if p.tonGiao and p.tonGiao != 'Không':
                        col7_parts.append(p.tonGiao)
                    cells[6].text = "\n".join(col7_parts)
                    cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 8: Văn hóa
                    cells[7].text = p.trinhDoVanHoa or ''
                    cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 9: Qua trường / Ngành học / Cấp học / Thời gian
                    col9_parts = []
                    if p.quaTruong:
                        col9_parts.append(p.quaTruong)
                    if p.nganhHoc:
                        col9_parts.append(p.nganhHoc)
                    if p.capHoc:
                        col9_parts.append(p.capHoc)
                    if p.thoiGianDaoTao:
                        col9_parts.append(p.thoiGianDaoTao)
                    cells[8].text = "\n".join(col9_parts)
                    cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 10: Quê quán / Trú quán / Khi cần báo tin cho ai SĐT
                    col10_parts = []
                    if p.queQuan:
                        col10_parts.append(p.queQuan)
                    if p.truQuan:
                        col10_parts.append(p.truQuan)
                    if p.lienHeKhiCan or p.soDienThoaiLienHe:
                        lien_he = p.lienHeKhiCan or ''
                        if p.soDienThoaiLienHe:
                            if lien_he:
                                lien_he += f" {p.soDienThoaiLienHe}"
                            else:
                                lien_he = p.soDienThoaiLienHe
                        col10_parts.append(lien_he)
                    cells[9].text = "\n".join(col10_parts)
                    cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 11: Họ tên cha / Họ tên mẹ / Họ tên vợ (kèm SĐT)
                    col11_parts = []
                    ho_ten_cha = p.hoTenCha or ''
                    ho_ten_me = p.hoTenMe or ''
                    ho_ten_vo = p.hoTenVo or ''
                    sdt_cha = ''
                    sdt_me = ''
                    sdt_vo = ''
                    
                    # Lấy từ bảng nguoi_than nếu chưa có
                    if db_service:
                        try:
                            nguoi_than_list = db_service.get_nguoi_than_by_personnel(p.id)
                            for nguoi_than in nguoi_than_list:
                                if not nguoi_than.hoTen:
                                    continue
                                moi_quan_he = (nguoi_than.moiQuanHe or '').lower()
                                if 'bố' in moi_quan_he or 'cha' in moi_quan_he:
                                    if not ho_ten_cha:
                                        ho_ten_cha = nguoi_than.hoTen
                                    if nguoi_than.soDienThoai:
                                        sdt_cha = nguoi_than.soDienThoai
                                elif 'mẹ' in moi_quan_he or 'me' in moi_quan_he:
                                    if not ho_ten_me:
                                        ho_ten_me = nguoi_than.hoTen
                                    if nguoi_than.soDienThoai:
                                        sdt_me = nguoi_than.soDienThoai
                                elif 'vợ' in moi_quan_he or 'vo' in moi_quan_he or 'chồng' in moi_quan_he or 'chong' in moi_quan_he:
                                    if not ho_ten_vo:
                                        ho_ten_vo = nguoi_than.hoTen
                                    if nguoi_than.soDienThoai:
                                        sdt_vo = nguoi_than.soDienThoai
                        except:
                            pass
                    
                    # Thêm vào danh sách theo thứ tự: cha, mẹ, vợ (kèm SĐT nếu có)
                    if ho_ten_cha:
                        if sdt_cha:
                            col11_parts.append(f"{ho_ten_cha} - {sdt_cha}")
                        else:
                            col11_parts.append(ho_ten_cha)
                    if ho_ten_me:
                        if sdt_me:
                            col11_parts.append(f"{ho_ten_me} - {sdt_me}")
                        else:
                            col11_parts.append(ho_ten_me)
                    if ho_ten_vo:
                        if sdt_vo:
                            col11_parts.append(f"{ho_ten_vo} - {sdt_vo}")
                        else:
                            col11_parts.append(ho_ten_vo)
                    
                    cells[10].text = "\n".join(col11_parts)
                    cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 12: Đơn vị đang làm nhiệm vụ (lấy từ p.unitId, hiển thị cả đơn vị cha và con)
                    ten_don_vi = ''
                    if p.unitId and db_service:
                        try:
                            current_unit = db_service.get_unit_by_id(p.unitId)
                            if current_unit:
                                # Lấy tên đơn vị con
                                ten_don_vi_con = current_unit.ten
                                
                                # Lấy đơn vị cha nếu có
                                if current_unit.parentId:
                                    parent_unit = db_service.get_unit_by_id(current_unit.parentId)
                                    if parent_unit:
                                        ten_don_vi = f"{parent_unit.ten} / {ten_don_vi_con}"
                                    else:
                                        ten_don_vi = ten_don_vi_con
                                else:
                                    ten_don_vi = ten_don_vi_con
                        except:
                            pass
                    # Nếu không có unitId, dùng unit từ units_data (sub-header)
                    if not ten_don_vi and unit:
                        ten_don_vi = unit.ten
                    cells[11].text = ten_don_vi
                    cells[11].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Cột 13: Ghi chú
                    cells[12].text = p.ghiChu or ''
                    cells[12].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Nếu không có units_data, xuất như cũ (tất cả quân nhân trong một danh sách)
            for idx, p in enumerate(personnel_list, 1):
                row = table.add_row()
                cells = row.cells
                
                # Format tất cả cells
                for cell in cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    
                    # Shading - nền trắng
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'FFFFFF')
                    shading_elm.set(qn('w:val'), 'clear')
                    tcPr.append(shading_elm)
                    
                    # Borders cho từng cell
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), '000000')
                        tcPr.append(border)
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                
                # Cột 1: STT
                cells[0].text = str(idx)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Cột 2: Họ và tên khai sinh / Họ và tên thường dùng
                ho_ten = p.hoTen or ''
                if p.hoTenThuongDung:
                    ho_ten += f"\n{p.hoTenThuongDung}"
                cells[1].text = ho_ten
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 3: Ngày tháng năm sinh / Cấp bậc / Ngày nhận
                col3_parts = []
                if p.ngaySinh:
                    col3_parts.append(p.ngaySinh)
                if p.capBac:
                    col3_parts.append(p.capBac)
                if p.ngayNhanCapBac:
                    # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                    ngay_nhan = p.ngayNhanCapBac
                    if '/' in ngay_nhan:
                        parts = ngay_nhan.split('/')
                        if len(parts) == 3:
                            ngay_nhan = f"{parts[1]}/{parts[2]}"
                    col3_parts.append(ngay_nhan)
                cells[2].text = "\n".join(col3_parts)
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 4: Chức vụ / Ngày nhận
                col4_parts = []
                if p.chucVu:
                    col4_parts.append(p.chucVu)
                if p.ngayNhanChucVu:
                    # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                    ngay_nhan = p.ngayNhanChucVu
                    if '/' in ngay_nhan:
                        parts = ngay_nhan.split('/')
                        if len(parts) == 3:
                            ngay_nhan = f"{parts[1]}/{parts[2]}"
                    col4_parts.append(ngay_nhan)
                cells[3].text = "\n".join(col4_parts)
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 5: Nhập ngũ / Xuất ngũ
                col5_parts = []
                if p.nhapNgu:
                    # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                    nhap_ngu = p.nhapNgu
                    if '/' in nhap_ngu:
                        parts = nhap_ngu.split('/')
                        if len(parts) == 3:
                            nhap_ngu = f"{parts[1]}/{parts[2]}"
                    col5_parts.append(nhap_ngu)
                if p.xuatNgu:
                    # Rút ngắn: DD/MM/YYYY -> MM/YYYY
                    xuat_ngu = p.xuatNgu
                    if '/' in xuat_ngu:
                        parts = xuat_ngu.split('/')
                        if len(parts) == 3:
                            xuat_ngu = f"{parts[1]}/{parts[2]}"
                    col5_parts.append(xuat_ngu)
                cells[4].text = "\n".join(col5_parts)
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 6: N. vào đoàn / N.vào đảng Chính thức
                col6_parts = []
                if p.thongTinKhac.doan.ngayVao:
                    # Giữ nguyên format DD/MM/YYYY
                    col6_parts.append(p.thongTinKhac.doan.ngayVao)
                if p.thongTinKhac.dang.ngayVao:
                    # Giữ nguyên format DD/MM/YYYY
                    col6_parts.append(p.thongTinKhac.dang.ngayVao)
                if p.thongTinKhac.dang.ngayChinhThuc:
                    # Giữ nguyên format DD/MM/YYYY
                    col6_parts.append(p.thongTinKhac.dang.ngayChinhThuc)
                cells[5].text = "\n".join(col6_parts)
                cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 7: Thành phần GĐ / Dân tộc / Tôn giáo
                col7_parts = []
                if p.thanhPhanGiaDinh:
                    col7_parts.append(p.thanhPhanGiaDinh)
                if p.danToc:
                    col7_parts.append(p.danToc)
                if p.tonGiao and p.tonGiao != 'Không':
                    col7_parts.append(p.tonGiao)
                cells[6].text = "\n".join(col7_parts)
                cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 8: Văn hóa
                cells[7].text = p.trinhDoVanHoa or ''
                cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 9: Qua trường / Ngành học / Cấp học / Thời gian
                col9_parts = []
                if p.quaTruong:
                    col9_parts.append(p.quaTruong)
                if p.nganhHoc:
                    col9_parts.append(p.nganhHoc)
                if p.capHoc:
                    col9_parts.append(p.capHoc)
                if p.thoiGianDaoTao:
                    col9_parts.append(p.thoiGianDaoTao)
                cells[8].text = "\n".join(col9_parts)
                cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 10: Quê quán / Trú quán / Khi cần báo tin cho ai SĐT
                col10_parts = []
                if p.queQuan:
                    col10_parts.append(p.queQuan)
                if p.truQuan:
                    col10_parts.append(p.truQuan)
                if p.lienHeKhiCan or p.soDienThoaiLienHe:
                    lien_he = p.lienHeKhiCan or ''
                    if p.soDienThoaiLienHe:
                        if lien_he:
                            lien_he += f" {p.soDienThoaiLienHe}"
                        else:
                            lien_he = p.soDienThoaiLienHe
                    col10_parts.append(lien_he)
                cells[9].text = "\n".join(col10_parts)
                cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 11: Họ tên cha / Họ tên mẹ / Họ tên vợ (kèm SĐT)
                col11_parts = []
                ho_ten_cha = p.hoTenCha or ''
                ho_ten_me = p.hoTenMe or ''
                ho_ten_vo = p.hoTenVo or ''
                sdt_cha = ''
                sdt_me = ''
                sdt_vo = ''
                
                # Lấy từ bảng nguoi_than nếu chưa có
                if db_service:
                    try:
                        nguoi_than_list = db_service.get_nguoi_than_by_personnel(p.id)
                        for nguoi_than in nguoi_than_list:
                            if not nguoi_than.hoTen:
                                continue
                            moi_quan_he = (nguoi_than.moiQuanHe or '').lower()
                            if 'bố' in moi_quan_he or 'cha' in moi_quan_he:
                                if not ho_ten_cha:
                                    ho_ten_cha = nguoi_than.hoTen
                                if nguoi_than.soDienThoai:
                                    sdt_cha = nguoi_than.soDienThoai
                            elif 'mẹ' in moi_quan_he or 'me' in moi_quan_he:
                                if not ho_ten_me:
                                    ho_ten_me = nguoi_than.hoTen
                                if nguoi_than.soDienThoai:
                                    sdt_me = nguoi_than.soDienThoai
                            elif 'vợ' in moi_quan_he or 'vo' in moi_quan_he or 'chồng' in moi_quan_he or 'chong' in moi_quan_he:
                                if not ho_ten_vo:
                                    ho_ten_vo = nguoi_than.hoTen
                                if nguoi_than.soDienThoai:
                                    sdt_vo = nguoi_than.soDienThoai
                    except:
                        pass
                
                # Thêm vào danh sách theo thứ tự: cha, mẹ, vợ (kèm SĐT nếu có)
                if ho_ten_cha:
                    if sdt_cha:
                        col11_parts.append(f"{ho_ten_cha} - {sdt_cha}")
                    else:
                        col11_parts.append(ho_ten_cha)
                if ho_ten_me:
                    if sdt_me:
                        col11_parts.append(f"{ho_ten_me} - {sdt_me}")
                    else:
                        col11_parts.append(ho_ten_me)
                if ho_ten_vo:
                    if sdt_vo:
                        col11_parts.append(f"{ho_ten_vo} - {sdt_vo}")
                    else:
                        col11_parts.append(ho_ten_vo)
                
                cells[10].text = "\n".join(col11_parts)
                cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 12: Đơn vị đang làm nhiệm vụ (hiển thị cả đơn vị cha và con)
                ten_don_vi = ''
                if p.unitId and db_service:
                    try:
                        current_unit = db_service.get_unit_by_id(p.unitId)
                        if current_unit:
                            # Lấy tên đơn vị con
                            ten_don_vi_con = current_unit.ten
                            
                            # Lấy đơn vị cha nếu có
                            if current_unit.parentId:
                                parent_unit = db_service.get_unit_by_id(current_unit.parentId)
                                if parent_unit:
                                    ten_don_vi = f"{parent_unit.ten} / {ten_don_vi_con}"
                                else:
                                    ten_don_vi = ten_don_vi_con
                            else:
                                ten_don_vi = ten_don_vi_con
                    except:
                        pass
                cells[11].text = ten_don_vi
                cells[11].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Cột 13: Ghi chú
                cells[12].text = p.ghiChu or ''
                cells[12].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Thêm phần "Xe ct (xe 694)" ở cuối (nếu có)
        # Có thể lấy từ ghi chú hoặc thêm field riêng
        doc.add_paragraph()  # Khoảng trống
        
        # Lưu vào buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("Cần cài đặt python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

