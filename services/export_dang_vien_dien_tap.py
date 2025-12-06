"""
Hàm xuất Word cho Đảng viên tham gia diễn tập năm 2025
"""
import io
import sys
from typing import List
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from models.personnel import Personnel


def to_word_docx_dang_vien_dien_tap(personnel_list: List[Personnel],
                                     tieu_doan: str = "TIỂU ĐOÀN 38",
                                     dai_doi: str = "ĐẠI ĐỘI 3",
                                     dia_diem: str = "Đăk Lăk",
                                     nam: str = "2025",
                                     db_service=None) -> bytes:
    """
    Xuất danh sách Đảng viên tham gia diễn tập năm 2025 ra file Word
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Thiết lập trang Letter Landscape theo thông số từ Page Setup
        section = doc.sections[0]
        # Letter size: Width 27.94 cm, Height 21.59 cm (Landscape)
        section.page_width = Inches(11.0)  # 27.94 cm = 11 inches
        section.page_height = Inches(8.5)  # 21.59 cm = 8.5 inches
        # Margins: Top 2 cm, Bottom 1 cm, Left 1.27 cm, Right 0.95 cm
        section.top_margin = Inches(0.787)  # 2 cm
        section.bottom_margin = Inches(0.394)  # 1 cm
        section.left_margin = Inches(0.5)  # 1.27 cm
        section.right_margin = Inches(0.374)  # 0.95 cm
        # Header and Footer: 1.27 cm
        section.header_distance = Inches(0.5)  # 1.27 cm
        section.footer_distance = Inches(0.5)  # 1.27 cm
        
        # Font mặc định
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Header với 2 cột (trái và phải)
        header_table = doc.add_table(rows=2, cols=2)
        header_table.autofit = False
        
        # Cột trái: ĐẢNG BỘ TIỂU ĐOÀN 38 CHI BỘ ĐẠI ĐỘI 3
        left_cell = header_table.rows[0].cells[0]
        left_cell.paragraphs[0].clear()
        left_run1 = left_cell.paragraphs[0].add_run(f"ĐẢNG BỘ {tieu_doan}")
        left_run1.font.size = Pt(11)
        left_run1.font.name = 'Times New Roman'
        left_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        left_cell.paragraphs[0].add_run("\n")
        left_run2 = left_cell.paragraphs[0].add_run(f"CHI BỘ {dai_doi.upper()}")
        left_run2.font.size = Pt(11)
        left_run2.font.name = 'Times New Roman'
        left_run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Cột phải: ĐẢNG CỘNG SẢN VIỆT NAM
        right_cell = header_table.rows[0].cells[1]
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].clear()
        right_run1 = right_cell.paragraphs[0].add_run("ĐẢNG CỘNG SẢN VIỆT NAM")
        right_run1.font.size = Pt(11)
        right_run1.font.name = 'Times New Roman'
        right_run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Merge cells cho dòng 2 (date line)
        header_table.rows[1].cells[0].merge(header_table.rows[1].cells[1])
        date_cell = header_table.rows[1].cells[0]
        date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_cell.paragraphs[0].clear()
        
        date_text = f"{dia_diem}, ngày tháng năm {nam}"
        date_run = date_cell.paragraphs[0].add_run(date_text)
        date_run.font.size = Pt(11)
        date_run.font.name = 'Times New Roman'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Bỏ borders cho header table
        for row in header_table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tiêu đề: DANH SÁCH
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DANH SÁCH")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Subtitle: đảng viên tham gia diễn tập năm 2025
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"đảng viên tham gia diễn tập năm {nam}")
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        subtitle_run.underline = True
        
        doc.add_paragraph()  # Khoảng trống
        
        # Tạo bảng với 10 cột (gộp Họ Và Tên và Ngày sinh thành 1 cột)
        table = doc.add_table(rows=1, cols=10)
        table.style = None
        
        # Thiết lập borders
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table width (tính theo margins mới)
        # Page width: 11 inches, Left margin: 0.5 inches, Right margin: 0.374 inches
        table_width_inches = 11.0 - 0.5 - 0.374
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(table_width_inches * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        
        # Header row
        header_row = table.rows[0]
        headers = ['TT', 'Họ Và Tên\nNgày, tháng, năm sinh', 'CB\nCV', 'Đơn Vị', 
                   'Văn Hóa', 'Dân\nTộc', 'Tôn\nGiáo', 'CV\nĐảng', 'Quê quán\nTrú quán', 'Ghi chú']
        
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Xử lý header có xuống dòng
            if '\n' in header_text:
                parts = header_text.split('\n')
                for i, part in enumerate(parts):
                    if i > 0:
                        cell.paragraphs[0].add_run('\n')
                    run = cell.paragraphs[0].add_run(part)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            else:
                run = cell.paragraphs[0].add_run(header_text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Set cell width
            if idx == 0:  # TT
                cell.width = Inches(0.4)
            elif idx == 1:  # Họ Và Tên / Ngày sinh (gộp)
                cell.width = Inches(1.5)
            elif idx == 2:  # CB CV
                cell.width = Inches(0.6)
            elif idx == 3:  # Đơn Vị
                cell.width = Inches(0.5)
            elif idx == 4:  # Văn Hóa
                cell.width = Inches(0.5)
            elif idx == 5:  # Dân Tộc
                cell.width = Inches(0.5)
            elif idx == 6:  # Tôn Giáo
                cell.width = Inches(0.5)
            elif idx == 7:  # CV Đảng
                cell.width = Inches(0.6)
            elif idx == 8:  # Quê quán Trú quán
                cell.width = Inches(1.5)
            else:  # Ghi chú
                cell.width = Inches(1.0)
        
        # Data rows
        for idx, person in enumerate(personnel_list, 1):
            row = table.add_row()
            
            # TT
            cell = row.cells[0]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(idx))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Họ Và Tên / Ngày, tháng, năm sinh (gộp thành 1 cột)
            cell = row.cells[1]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            ho_ten = person.hoTen or ''
            ngay_sinh = person.ngaySinh or ''
            ho_ten_ngay_sinh = f"{ho_ten}\n{ngay_sinh}".strip('\n')
            run = cell.paragraphs[0].add_run(ho_ten_ngay_sinh)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # CB CV
            cell = row.cells[2]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cb_cv = f"{person.capBac or ''}/{person.chucVu or ''}".strip('/')
            run = cell.paragraphs[0].add_run(cb_cv)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Đơn Vị
            cell = row.cells[3]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(person.donVi or '')
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Văn Hóa
            cell = row.cells[4]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(person.trinhDoVanHoa or '')
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Dân Tộc
            cell = row.cells[5]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(person.danToc or '')
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Tôn Giáo
            cell = row.cells[6]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            ton_giao = person.tonGiao or 'Không'
            run = cell.paragraphs[0].add_run(ton_giao)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # CV Đảng
            cell = row.cells[7]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(person.thongTinKhac.dang.chucVuDang or '')
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Quê quán Trú quán
            cell = row.cells[8]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            que_quan = person.queQuan or ''
            tru_quan = person.truQuan or ''
            que_tru = f"{que_quan} / {tru_quan}".strip(' / ').strip()
            run = cell.paragraphs[0].add_run(que_tru)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            # Ghi chú - lấy từ ghi chú riêng của tab
            cell = row.cells[9]
            cell.paragraphs[0].clear()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Lấy ghi chú riêng từ database
            ghi_chu = ''
            if db_service:
                ghi_chu = db_service.get_dang_vien_dien_tap_ghi_chu(person.id)
            run = cell.paragraphs[0].add_run(ghi_chu)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        doc.add_paragraph()  # Khoảng trống
        
        # Footer: T/M CHI BỘ BÍ THƯ
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("T/M CHI BỘ")
        footer_run.font.size = Pt(11)
        footer_run.font.name = 'Times New Roman'
        footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer2_run = footer2.add_run("BÍ THƯ")
        footer2_run.bold = True
        footer2_run.font.size = Pt(11)
        footer2_run.font.name = 'Times New Roman'
        footer2_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        raise Exception(f"Lỗi khi xuất file Word: {str(e)}")

